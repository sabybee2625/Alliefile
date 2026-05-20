from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, status, Form, Query, Body, Request, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from fastapi.responses import FileResponse, StreamingResponse, Response
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr, validator
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import secrets
import jwt
from passlib.context import CryptContext
import aiofiles
import json
import base64
import io
import zipfile
import subprocess
import tempfile
import hashlib
import asyncio

# Load environment before importing config
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Import our modules
from config import config, get_plan_limits, PLAN_LIMITS
from rate_limiter import (
    rate_limiter, 
    check_rate_limit_login, 
    check_rate_limit_register,
    check_rate_limit_analysis,
    check_rate_limit_assistant,
    get_client_ip
)
from security import (
    SecurityHeadersMiddleware,
    ErrorHandlingMiddleware,
    AccessLogMiddleware,
    log_share_access,
    validate_user_owns_resource
)
from storage import get_storage_backend, compute_file_hash, LocalStorage, GridFSStorage, EmergentObjectStorage
from emailing import send_welcome_email_background, send_password_reset_email_background
from admin import register_admin_routes
from piece_classifier import classify_piece

# MongoDB connection with SSL certificate handling for Atlas
import certifi

# Use certifi only for Atlas connections (mongodb+srv)
if config.MONGO_URL.startswith("mongodb+srv"):
    client = AsyncIOMotorClient(config.MONGO_URL, tlsCAFile=certifi.where())
else:
    client = AsyncIOMotorClient(config.MONGO_URL)
db = client[config.DB_NAME]

# Create the main app with production settings
app = FastAPI(
    title="AlliéFile — Votre allié juridique intelligent",
    description="SaaS sécurisé pour la constitution et l'analyse de dossiers juridiques",
    version="1.1.0",
    docs_url="/api/docs" if config.DEBUG else None,  # Disable docs in production
    redoc_url="/api/redoc" if config.DEBUG else None,
)

# Create router with /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Storage backend (abstracted for cloud migration)
storage = get_storage_backend(db_instance=db)

# Logging
log_level = logging.DEBUG if config.DEBUG else logging.INFO
logging.basicConfig(
    level=log_level, 
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Log startup info
logger.info(f"Starting application in {config.ENV.value} mode")
if config.IS_PRODUCTION:
    logger.info("Production mode: API docs disabled, security headers enabled")

# Analysis queue lock (simple in-memory for now)
analysis_locks = {}

# Analysis rate limiting constants
MAX_CONCURRENT_ANALYSES = config.MAX_CONCURRENT_ANALYSES
ANALYSIS_RATE_LIMIT_SECONDS = 2  # Minimum seconds between analyses

# ===================== MODELS =====================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str
    
    @validator('password')
    def password_strength(cls, v):
        if len(v) < 8:
            raise ValueError('Le mot de passe doit contenir au moins 8 caractères')
        return v

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    id: str
    email: str
    name: str
    plan: str = "free"
    plan_expires_at: Optional[str] = None
    created_at: str

class UserStats(BaseModel):
    total_dossiers: int
    total_pieces: int
    pieces_ready: int
    pieces_to_verify: int
    pieces_by_type: Dict[str, int]
    active_share_links: int
    storage_used_mb: float
    plan: str
    plan_limits: Dict[str, Any]
    assistant_uses_today: int

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class DossierCreate(BaseModel):
    title: str
    description: Optional[str] = ""

class DossierUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None

class DossierResponse(BaseModel):
    id: str
    title: str
    description: str
    user_id: str
    created_at: str
    updated_at: str
    piece_count: int = 0

class AIProposal(BaseModel):
    # Champs rendus optionnels pour rester compatibles avec
    # les propositions partielles écrites par le classifieur déterministe
    # (qui n'enrichit que tags_thematiques/sujets_concernes/nature_document).
    type_piece: Optional[str] = None
    type_confidence: Optional[str] = None
    date_document: Optional[str] = None
    date_confidence: str = "faible"
    titre: Optional[str] = None
    titre_confidence: str = "moyen"
    resume_qui: Optional[str] = None
    resume_quoi: Optional[str] = None
    resume_ou: Optional[str] = None
    resume_element_cle: Optional[str] = None
    mots_cles: List[str] = []
    extrait_justificatif: Optional[str] = None
    # V2 enrichment (neutral classification)
    tags_thematiques: List[str] = []   # ex: famille, travail, sante, finances, logement, violence, harcelement, administratif
    sujets_concernes: List[str] = []   # ex: utilisateur, conjoint, enfant, employeur, bailleur, tiers, administration
    nature_document: Optional[str] = None  # officiel | prive | temoignage | medical | financier | autre
    # V3 — taxonomie 3 niveaux
    sous_domaine: Optional[str] = None      # ex: "Violence conjugale", "Litige locatif"
    type_specifique: Optional[str] = None   # ex: "psychologique", "loyers"
    source_qualifiee: Optional[str] = None  # "PRO" | "PRIVÉ" | None

class PieceValidation(BaseModel):
    type_piece: str
    date_document: Optional[str] = None
    titre: str
    resume_qui: Optional[str] = None
    resume_quoi: Optional[str] = None
    resume_ou: Optional[str] = None
    resume_element_cle: Optional[str] = None
    mots_cles: List[str] = []
    tags_thematiques: List[str] = []
    sujets_concernes: List[str] = []
    nature_document: Optional[str] = None
    sous_domaine: Optional[str] = None
    type_specifique: Optional[str] = None
    source_qualifiee: Optional[str] = None

class PieceResponse(BaseModel):
    id: str
    dossier_id: str
    numero: int
    filename: str
    original_filename: str
    file_type: str
    file_size: int = 0
    file_hash: Optional[str] = None
    is_duplicate: bool = False
    file_missing: bool = False
    source: str = "upload"  # 'upload' or 'camera'
    status: str
    analysis_status: str = "pending"  # pending, queued, analyzing, complete, error
    analysis_error: Optional[str] = None
    analysis_queued_at: Optional[str] = None
    analysis_started_at: Optional[str] = None
    analysis_completed_at: Optional[str] = None
    ai_proposal: Optional[AIProposal] = None
    validated_data: Optional[PieceValidation] = None
    created_at: str
    updated_at: str

class ShareLinkCreate(BaseModel):
    expires_in_days: int = 7
    piece_ids: Optional[List[str]] = None  # None = all pieces, [] = none, [...] = specific

class ShareLinkResponse(BaseModel):
    id: str
    dossier_id: str
    token: str
    expires_at: str
    created_at: str

class AssistantRequest(BaseModel):
    document_type: str
    jurisdiction: Optional[str] = None
    piece_ids: List[str] = []
    date_start: Optional[str] = None
    date_end: Optional[str] = None

class AssistantResponse(BaseModel):
    content: str
    pieces_used: List[int]
    warnings: List[str] = []

class DeleteManyRequest(BaseModel):
    piece_ids: List[str]

class DuplicateCheckResponse(BaseModel):
    is_duplicate: bool
    existing_piece_id: Optional[str] = None
    existing_piece_numero: Optional[int] = None
    existing_filename: Optional[str] = None

class DuplicateErrorDetail(BaseModel):
    message: str
    existing_piece_id: str
    existing_piece_numero: int
    existing_filename: str

class QueueAnalysisRequest(BaseModel):
    piece_ids: List[str] = []  # Empty = all pending pieces

# Share link models for advanced sharing
class ShareLinkCreateAdvanced(BaseModel):
    expires_in_days: int = 7
    piece_ids: Optional[List[str]] = None  # None = all pieces
    filter_type: Optional[str] = None  # Filter by piece type
    filter_status: Optional[str] = None  # Filter by status
    filter_keywords: Optional[List[str]] = None
    password: Optional[str] = None  # Optional password protection

class PromoCodeCreate(BaseModel):
    code: str
    discount_percent: Optional[int] = None
    discount_amount: Optional[float] = None
    max_uses: int = -1  # -1 for unlimited
    expires_at: Optional[str] = None
    plan_restriction: Optional[str] = None  # Only for specific plan

# ===================== AUTH HELPERS =====================

def hash_password(password: str) -> str:
    return pwd_context.hash(password)

def verify_password(plain: str, hashed: str) -> bool:
    return pwd_context.verify(plain, hashed)

def create_token(user_id: str) -> str:
    payload = {
        "sub": user_id,
        "exp": datetime.now(timezone.utc) + timedelta(hours=config.JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, config.JWT_SECRET, algorithm=config.JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, config.JWT_SECRET, algorithms=[config.JWT_ALGORITHM])
        user_id = payload.get("sub")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        user = await db.users.find_one({"id": user_id}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

async def get_user_plan(user: dict) -> str:
    """Get user's current plan, checking expiration"""
    plan = user.get("plan", "free")
    expires_at = user.get("plan_expires_at") or user.get("current_period_end")
    plan_status = user.get("plan_status", "active" if plan != "free" else None)
    
    if plan != "free" and expires_at:
        if datetime.fromisoformat(expires_at) < datetime.now(timezone.utc):
            # Plan expired, revert to free
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {
                    "plan": "free",
                    "plan_status": "expired",
                    "plan_expires_at": None,
                    "current_period_end": None
                }}
            )
            logger.info(f"User {user['id']} plan expired, reverted to free")
            return "free"
    
    # Handle canceled subscriptions that reach end of period
    if plan_status == "canceled" and expires_at:
        if datetime.fromisoformat(expires_at) < datetime.now(timezone.utc):
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {
                    "plan": "free",
                    "plan_status": "expired",
                    "plan_expires_at": None,
                    "current_period_end": None
                }}
            )
            logger.info(f"User {user['id']} canceled plan ended, reverted to free")
            return "free"
    
    return plan

async def check_plan_limit(user: dict, limit_type: str, current_count: int = 0) -> bool:
    """
    Check if user is within their plan limits
    Returns True if within limits, raises HTTPException if exceeded
    """
    plan = await get_user_plan(user)
    limits = get_plan_limits(plan)
    
    limit_value = getattr(limits, limit_type, -1)
    
    if limit_value == -1:  # Unlimited
        return True
    
    # Map limit types to French labels
    limit_labels = {
        "max_dossiers": f"nombre de dossiers ({limit_value} max)",
        "max_total_pieces": f"nombre total de pièces ({limit_value} max)",
        "max_pieces_per_dossier": f"pièces par dossier ({limit_value} max)",
        "max_share_links": f"liens de partage ({limit_value} max)",
        "assistant_daily_limit": f"utilisations de l'assistant par jour ({limit_value} max)"
    }
    
    if current_count >= limit_value:
        label = limit_labels.get(limit_type, limit_type)
        raise HTTPException(
            status_code=403,
            detail={
                "error": "PLAN_LIMIT_EXCEEDED",
                "message": f"Limite atteinte : {label}. Passez au plan Essentiel pour continuer.",
                "plan": plan,
                "limit_type": limit_type,
                "current": current_count,
                "max": limit_value,
                "upgrade_url": "/pricing"
            }
        )
    
    return True

# ===================== AUTH ROUTES =====================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(data: UserCreate, request: Request, background_tasks: BackgroundTasks):
    # Rate limiting
    client_ip = get_client_ip(request)
    await check_rate_limit_register(request, client_ip)
    
    existing = await db.users.find_one({"email": data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    user_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    user_doc = {
        "id": user_id,
        "email": data.email,
        "name": data.name,
        "password_hash": hash_password(data.password),
        "plan": "free",
        "plan_expires_at": None,
        "stripe_customer_id": None,
        "assistant_uses_today": 0,
        "assistant_last_reset": now,
        "created_at": now
    }
    await db.users.insert_one(user_doc)
    
    # Fire-and-forget welcome email (non-blocking, never fails the request)
    try:
        background_tasks.add_task(send_welcome_email_background, data.email, data.name)
    except Exception as e:
        logger.error(f"Failed to enqueue welcome email for {data.email}: {e}")
    
    token = create_token(user_id)
    return TokenResponse(
        access_token=token,
        user=UserResponse(id=user_id, email=data.email, name=data.name, plan="free", created_at=now)
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(data: UserLogin, request: Request):
    # Rate limiting
    client_ip = get_client_ip(request)
    await check_rate_limit_login(request, client_ip)
    
    user = await db.users.find_one({"email": data.email}, {"_id": 0})
    if not user or not verify_password(data.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    # Check if account has scheduled deletion - cancel it on login
    if user.get("scheduled_deletion"):
        await db.users.update_one(
            {"id": user["id"]},
            {"$unset": {"scheduled_deletion": ""}}
        )
        logger.info(f"Account deletion cancelled on login: user={user['id']}")
    
    token = create_token(user["id"])
    return TokenResponse(
        access_token=token,
        user=UserResponse(
            id=user["id"], 
            email=user["email"], 
            name=user["name"], 
            plan=user.get("plan", "free"),
            plan_expires_at=user.get("plan_expires_at"),
            created_at=user["created_at"]
        )
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(user: dict = Depends(get_current_user)):
    return UserResponse(
        id=user["id"], 
        email=user["email"], 
        name=user["name"], 
        plan=user.get("plan", "free"),
        plan_expires_at=user.get("plan_expires_at"),
        created_at=user["created_at"]
    )


# ============================================================
# PASSWORD RESET
# ============================================================

class PasswordResetRequest(BaseModel):
    email: EmailStr

class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str


@api_router.post("/auth/password-reset/request")
async def request_password_reset(
    data: PasswordResetRequest,
    request: Request,
    background_tasks: BackgroundTasks,
):
    """Request a password reset. Always returns 200 to avoid email enumeration."""
    # Rate limiting (reuse login limiter for security)
    client_ip = get_client_ip(request)
    try:
        await check_rate_limit_login(request, client_ip)
    except Exception:
        pass

    user = await db.users.find_one({"email": data.email.lower()}, {"_id": 0})
    if user:
        token = secrets.token_urlsafe(32)
        now = datetime.now(timezone.utc)
        expires = now + timedelta(hours=1)
        await db.password_reset_tokens.insert_one({
            "token": token,
            "user_id": user["id"],
            "email": user["email"],
            "created_at": now.isoformat(),
            "expires_at": expires.isoformat(),
            "used": False,
        })

        # Build reset URL — prefer APP_PUBLIC_URL, fallback to request origin
        base_url = os.environ.get("APP_PUBLIC_URL")
        if not base_url:
            origin = request.headers.get("origin") or request.headers.get("referer", "")
            base_url = origin.rstrip("/").rstrip("/login").rstrip("/register") if origin else "https://alliefile.com"
        reset_url = f"{base_url.rstrip('/')}/reset-password?token={token}"

        try:
            background_tasks.add_task(
                send_password_reset_email_background,
                user["email"],
                user.get("name", ""),
                reset_url,
            )
        except Exception as e:
            logger.error(f"Failed to enqueue reset email for {user['email']}: {e}")
        logger.info(f"Password reset requested for user={user['id']}")
    else:
        logger.info(f"Password reset requested for unknown email: {data.email}")

    return {"ok": True, "message": "Si un compte existe avec cet email, un lien de réinitialisation a été envoyé."}


@api_router.post("/auth/password-reset/confirm")
async def confirm_password_reset(data: PasswordResetConfirm):
    """Confirm password reset using the token."""
    if len(data.new_password) < 8:
        raise HTTPException(status_code=400, detail="Le mot de passe doit contenir au moins 8 caractères")

    token_doc = await db.password_reset_tokens.find_one({"token": data.token}, {"_id": 0})
    if not token_doc:
        raise HTTPException(status_code=400, detail="Lien invalide ou déjà utilisé")
    if token_doc.get("used"):
        raise HTTPException(status_code=400, detail="Ce lien a déjà été utilisé")
    expires_at = token_doc.get("expires_at", "")
    try:
        if datetime.fromisoformat(expires_at) < datetime.now(timezone.utc):
            raise HTTPException(status_code=400, detail="Lien expiré, redemandez une réinitialisation")
    except ValueError:
        raise HTTPException(status_code=400, detail="Lien invalide")

    # Update password
    new_hash = hash_password(data.new_password)
    result = await db.users.update_one(
        {"id": token_doc["user_id"]},
        {"$set": {"password_hash": new_hash, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    if result.matched_count == 0:
        raise HTTPException(status_code=400, detail="Compte introuvable")

    # Invalidate the token
    await db.password_reset_tokens.update_one(
        {"token": data.token},
        {"$set": {"used": True, "used_at": datetime.now(timezone.utc).isoformat()}}
    )
    logger.info(f"Password reset confirmed for user={token_doc['user_id']}")
    return {"ok": True, "message": "Mot de passe réinitialisé avec succès"}

@api_router.get("/auth/stats", response_model=UserStats)
async def get_user_stats(user: dict = Depends(get_current_user)):
    """Get user statistics and plan usage for dashboard"""
    user_id = user["id"]
    plan = await get_user_plan(user)
    limits = get_plan_limits(plan)
    
    # Count dossiers
    total_dossiers = await db.dossiers.count_documents({"user_id": user_id})
    
    # Count pieces
    pipeline = [
        {"$match": {"dossier_id": {"$in": [d["id"] async for d in db.dossiers.find({"user_id": user_id}, {"id": 1})]}}},
        {"$group": {
            "_id": None,
            "total": {"$sum": 1},
            "ready": {"$sum": {"$cond": [{"$eq": ["$status", "pret"]}, 1, 0]}},
            "to_verify": {"$sum": {"$cond": [{"$eq": ["$status", "a_verifier"]}, 1, 0]}},
            "total_size": {"$sum": "$file_size"}
        }}
    ]
    
    # Get dossier IDs first
    dossier_ids = [d["id"] async for d in db.dossiers.find({"user_id": user_id}, {"id": 1})]
    
    pieces_stats = await db.pieces.aggregate([
        {"$match": {"dossier_id": {"$in": dossier_ids}}},
        {"$group": {
            "_id": None,
            "total": {"$sum": 1},
            "ready": {"$sum": {"$cond": [{"$eq": ["$status", "pret"]}, 1, 0]}},
            "to_verify": {"$sum": {"$cond": [{"$eq": ["$status", "a_verifier"]}, 1, 0]}},
            "total_size": {"$sum": "$file_size"}
        }}
    ]).to_list(1)
    
    stats = pieces_stats[0] if pieces_stats else {"total": 0, "ready": 0, "to_verify": 0, "total_size": 0}
    
    # Pieces by type
    type_pipeline = [
        {"$match": {"dossier_id": {"$in": dossier_ids}}},
        {"$group": {"_id": "$file_type", "count": {"$sum": 1}}}
    ]
    type_results = await db.pieces.aggregate(type_pipeline).to_list(100)
    pieces_by_type = {r["_id"]: r["count"] for r in type_results if r["_id"]}
    
    # Active share links
    now = datetime.now(timezone.utc).isoformat()
    active_links = await db.share_links.count_documents({
        "dossier_id": {"$in": dossier_ids},
        "expires_at": {"$gt": now},
        "revoked": {"$ne": True}
    })
    
    # Assistant uses today
    assistant_uses = user.get("assistant_uses_today", 0)
    last_reset = user.get("assistant_last_reset")
    if last_reset:
        last_reset_date = datetime.fromisoformat(last_reset).date()
        if last_reset_date < datetime.now(timezone.utc).date():
            assistant_uses = 0
    
    return UserStats(
        total_dossiers=total_dossiers,
        total_pieces=stats["total"],
        pieces_ready=stats["ready"],
        pieces_to_verify=stats["to_verify"],
        pieces_by_type=pieces_by_type,
        active_share_links=active_links,
        storage_used_mb=stats["total_size"] / (1024 * 1024),
        plan=plan,
        plan_limits=limits.dict(),
        assistant_uses_today=assistant_uses
    )

# ===================== PAYMENT ROUTES =====================

from payments import (
    SUBSCRIPTION_PLANS, 
    CreateCheckoutRequest, 
    CheckoutResponse,
    get_plan_price,
    validate_promo_code,
    apply_promo_discount,
    increment_promo_usage,
    create_checkout_session,
    get_checkout_status,
    handle_stripe_webhook,
    STRIPE_AVAILABLE,
    normalize_plan_id
)

@api_router.get("/payments/plans")
async def get_subscription_plans():
    """Get available subscription plans"""
    return {
        "plans": SUBSCRIPTION_PLANS,
        "stripe_available": STRIPE_AVAILABLE
    }

@api_router.post("/payments/checkout", response_model=CheckoutResponse)
async def create_payment_checkout(
    data: CreateCheckoutRequest,
    request: Request,
    user: dict = Depends(get_current_user)
):
    """Create a Stripe checkout session for subscription"""
    if not STRIPE_AVAILABLE:
        raise HTTPException(status_code=503, detail="Paiements non disponibles actuellement")
    
    api_key = os.environ.get("STRIPE_API_KEY")
    if not api_key:
        raise HTTPException(status_code=503, detail="Stripe non configuré")
    
    # Validate plan (accept both "essentiel"/"pro" slugs and internal keys)
    data.plan_id = normalize_plan_id(data.plan_id)
    if data.plan_id not in SUBSCRIPTION_PLANS:
        raise HTTPException(status_code=400, detail="Plan invalide")
    
    # Get base price (SERVER-SIDE only)
    base_amount, currency = get_plan_price(data.plan_id, data.billing_period)
    
    # Validate promo code
    discount_info = {"valid": False}
    if data.promo_code:
        discount_info = await validate_promo_code(db, data.promo_code, data.plan_id)
        if not discount_info["valid"]:
            raise HTTPException(status_code=400, detail=discount_info.get("error", "Code promo invalide"))
    
    # Apply discount
    final_amount = await apply_promo_discount(base_amount, discount_info)
    discount_applied = base_amount - final_amount
    
    # Get origin URL from request
    origin_url = request.headers.get("origin") or str(request.base_url).rstrip("/")
    webhook_url = f"{str(request.base_url).rstrip('/')}/api/webhook/stripe"
    
    try:
        # Create Stripe checkout session
        session = await create_checkout_session(
            api_key=api_key,
            webhook_url=webhook_url,
            origin_url=origin_url,
            user_id=user["id"],
            user_email=user["email"],
            plan_id=data.plan_id,
            billing_period=data.billing_period,
            amount=final_amount,
            currency=currency,
            promo_code=data.promo_code
        )
        
        # Create payment transaction record BEFORE redirect
        now = datetime.now(timezone.utc).isoformat()
        transaction = {
            "id": str(uuid.uuid4()),
            "user_id": user["id"],
            "user_email": user["email"],
            "session_id": session.session_id,
            "plan_id": data.plan_id,
            "billing_period": data.billing_period,
            "amount": final_amount,
            "currency": currency,
            "promo_code": data.promo_code,
            "discount_applied": discount_applied,
            "status": "pending",
            "created_at": now,
            "updated_at": now
        }
        await db.payment_transactions.insert_one(transaction)
        
        logger.info(f"Created checkout session {session.session_id} for user {user['id']}")
        
        return CheckoutResponse(
            checkout_url=session.url,
            session_id=session.session_id
        )
    
    except Exception as e:
        logger.error(f"Error creating checkout session: {e}")
        raise HTTPException(status_code=500, detail="Erreur lors de la création du paiement")

@api_router.get("/payments/status/{session_id}")
async def check_payment_status(session_id: str, request: Request, user: dict = Depends(get_current_user)):
    """Check the status of a payment session and update user plan if paid"""
    if not STRIPE_AVAILABLE:
        raise HTTPException(status_code=503, detail="Paiements non disponibles")
    
    api_key = os.environ.get("STRIPE_API_KEY")
    if not api_key:
        raise HTTPException(status_code=503, detail="Stripe non configuré")
    
    webhook_url = f"{str(request.base_url).rstrip('/')}/api/webhook/stripe"
    
    # Get transaction record
    transaction = await db.payment_transactions.find_one(
        {"session_id": session_id, "user_id": user["id"]},
        {"_id": 0}
    )
    
    if not transaction:
        raise HTTPException(status_code=404, detail="Transaction non trouvée")
    
    # If already processed, return status
    if transaction["status"] == "paid":
        return {"status": "paid", "message": "Paiement déjà traité"}
    
    try:
        # Check with Stripe
        status = await get_checkout_status(api_key, webhook_url, session_id)
        
        now = datetime.now(timezone.utc).isoformat()
        
        if status.payment_status == "paid":
            # Update transaction
            await db.payment_transactions.update_one(
                {"session_id": session_id},
                {"$set": {"status": "paid", "updated_at": now}}
            )
            
            # Update user plan
            plan_id = transaction["plan_id"]
            billing_period = transaction["billing_period"]
            
            # Calculate expiration (current_period_end)
            if billing_period == "yearly":
                current_period_end = (datetime.now(timezone.utc) + timedelta(days=365)).isoformat()
            else:
                current_period_end = (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()
            
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {
                    "plan": plan_id,
                    "plan_status": "active",
                    "plan_expires_at": current_period_end,
                    "current_period_end": current_period_end,
                    "updated_at": now
                }}
            )
            
            # Increment promo code usage
            if transaction.get("promo_code"):
                await increment_promo_usage(db, transaction["promo_code"])
            
            logger.info(f"User {user['id']} upgraded to {plan_id}")
            
            return {
                "status": "paid",
                "plan": plan_id,
                "plan_status": "active",
                "current_period_end": current_period_end,
                "message": f"Bienvenue dans le plan {SUBSCRIPTION_PLANS[plan_id]['name']}!"
            }
        
        elif status.status == "expired":
            await db.payment_transactions.update_one(
                {"session_id": session_id},
                {"$set": {"status": "expired", "updated_at": now}}
            )
            return {"status": "expired", "message": "Session de paiement expirée"}
        
        else:
            return {"status": "pending", "message": "Paiement en cours de traitement"}
    
    except Exception as e:
        logger.error(f"Error checking payment status: {e}")
        raise HTTPException(status_code=500, detail="Erreur lors de la vérification du paiement")

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """
    Handle Stripe webhooks.
    Handles: checkout.session.completed, invoice.payment_succeeded,
             invoice.payment_failed, customer.subscription.deleted,
             customer.subscription.updated
    """
    if not STRIPE_AVAILABLE:
        return {"status": "ignored"}
    
    api_key = os.environ.get("STRIPE_API_KEY")
    webhook_url = f"{str(request.base_url).rstrip('/')}/api/webhook/stripe"
    
    body = await request.body()
    signature = request.headers.get("Stripe-Signature", "")
    
    try:
        event = await handle_stripe_webhook(api_key, webhook_url, body, signature)
        event_type = getattr(event, "event_type", None) or "unknown"
        logger.info(f"Stripe webhook received: {event_type}")
        now = datetime.now(timezone.utc).isoformat()

        # 1) Checkout finalisé => marquer transaction payée + upgrader plan
        if event_type == "checkout.session.completed":
            session_id = getattr(event, "session_id", None)
            if session_id:
                tx = await db.payment_transactions.find_one({"session_id": session_id}, {"_id": 0})
                if tx and tx.get("status") != "paid":
                    await db.payment_transactions.update_one(
                        {"session_id": session_id},
                        {"$set": {"status": "paid", "updated_at": now}}
                    )
                    # Upgrade user plan
                    plan_id = tx.get("plan_id")
                    billing_period = tx.get("billing_period", "monthly")
                    if plan_id and tx.get("user_id"):
                        days = 365 if billing_period == "yearly" else 30
                        expires = (datetime.now(timezone.utc) + timedelta(days=days)).isoformat()
                        await db.users.update_one(
                            {"id": tx["user_id"]},
                            {"$set": {
                                "plan": plan_id,
                                "plan_status": "active",
                                "plan_expires_at": expires,
                                "current_period_end": expires,
                                "updated_at": now,
                            }}
                        )
                        if tx.get("promo_code"):
                            await increment_promo_usage(db, tx["promo_code"])
                        logger.info(f"Webhook upgraded user {tx['user_id']} to {plan_id}")

        # 2) Facture payée => étendre la période
        elif event_type == "invoice.payment_succeeded":
            metadata = getattr(event, "metadata", {}) or {}
            user_id = metadata.get("user_id")
            if user_id:
                user = await db.users.find_one({"id": user_id}, {"_id": 0})
                if user:
                    billing = metadata.get("billing_period", "monthly")
                    days = 365 if billing == "yearly" else 30
                    expires = (datetime.now(timezone.utc) + timedelta(days=days)).isoformat()
                    await db.users.update_one(
                        {"id": user_id},
                        {"$set": {
                            "plan_status": "active",
                            "plan_expires_at": expires,
                            "current_period_end": expires,
                            "updated_at": now,
                        }}
                    )
                    logger.info(f"Webhook invoice.paid for user {user_id}, period extended")

        # 3) Échec de paiement => flag past_due
        elif event_type == "invoice.payment_failed":
            metadata = getattr(event, "metadata", {}) or {}
            user_id = metadata.get("user_id")
            if user_id:
                await db.users.update_one(
                    {"id": user_id},
                    {"$set": {"plan_status": "past_due", "updated_at": now}}
                )
                logger.warning(f"Webhook invoice.payment_failed for user {user_id}")

        # 4) Abonnement annulé => garder jusqu'à fin de période puis expirer
        elif event_type == "customer.subscription.deleted":
            metadata = getattr(event, "metadata", {}) or {}
            user_id = metadata.get("user_id")
            if user_id:
                await db.users.update_one(
                    {"id": user_id},
                    {"$set": {"plan_status": "canceled", "updated_at": now}}
                )
                logger.info(f"Webhook subscription.deleted for user {user_id}")

        # 5) Mise à jour d'abonnement (changement de plan, etc.)
        elif event_type == "customer.subscription.updated":
            metadata = getattr(event, "metadata", {}) or {}
            user_id = metadata.get("user_id")
            plan_id = metadata.get("plan_id")
            if user_id and plan_id:
                await db.users.update_one(
                    {"id": user_id},
                    {"$set": {"plan": plan_id, "plan_status": "active", "updated_at": now}}
                )
                logger.info(f"Webhook subscription.updated user={user_id} plan={plan_id}")

        return {"status": "received", "event_type": event_type}
    
    except Exception as e:
        logger.error(f"Webhook error: {e}", exc_info=True)
        # Return 200 to avoid Stripe retries for parse errors; log for investigation
        return {"status": "error", "error": str(e)}

@api_router.post("/payments/promo-codes", status_code=201)
async def create_promo_code(
    code: str = Body(...),
    discount_percent: Optional[int] = Body(None),
    discount_amount: Optional[float] = Body(None),
    max_uses: int = Body(-1),
    expires_at: Optional[str] = Body(None),
    plan_restriction: Optional[str] = Body(None),
    user: dict = Depends(get_current_user)
):
    """Create a promo code (admin only in production)"""
    # In production, add admin check here
    
    now = datetime.now(timezone.utc).isoformat()
    promo_doc = {
        "id": str(uuid.uuid4()),
        "code": code.upper(),
        "discount_percent": discount_percent,
        "discount_amount": discount_amount,
        "max_uses": max_uses,
        "uses": 0,
        "expires_at": expires_at,
        "plan_restriction": plan_restriction,
        "created_by": user["id"],
        "created_at": now
    }
    
    try:
        await db.promo_codes.insert_one(promo_doc)
        return {"message": "Code promo créé", "code": code.upper()}
    except Exception as e:
        if "duplicate" in str(e).lower():
            raise HTTPException(status_code=400, detail="Ce code existe déjà")
        raise HTTPException(status_code=500, detail="Erreur lors de la création")

@api_router.post("/payments/validate-promo")
async def validate_promo(
    code: str = Body(...),
    plan_id: str = Body(...),
    user: dict = Depends(get_current_user)
):
    """Validate a promo code before checkout"""
    plan_id = normalize_plan_id(plan_id)
    result = await validate_promo_code(db, code, plan_id)
    
    if not result["valid"]:
        raise HTTPException(status_code=400, detail=result.get("error", "Code invalide"))
    
    # Calculate discount preview
    base_amount, currency = get_plan_price(plan_id, "monthly")
    final_amount = await apply_promo_discount(base_amount, result)
    
    return {
        "valid": True,
        "discount_percent": result.get("discount_percent", 0),
        "discount_amount": result.get("discount_amount", 0),
        "original_price": base_amount,
        "final_price": final_amount,
        "currency": currency
    }

# ===================== DOSSIER ROUTES =====================

@api_router.post("/dossiers", response_model=DossierResponse)
async def create_dossier(data: DossierCreate, user: dict = Depends(get_current_user)):
    # Check plan limits
    current_count = await db.dossiers.count_documents({"user_id": user["id"]})
    await check_plan_limit(user, "max_dossiers", current_count)
    
    dossier_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    dossier_doc = {
        "id": dossier_id,
        "title": data.title,
        "description": data.description or "",
        "user_id": user["id"],
        "created_at": now,
        "updated_at": now
    }
    await db.dossiers.insert_one(dossier_doc)
    return DossierResponse(**dossier_doc, piece_count=0)

@api_router.get("/dossiers", response_model=List[DossierResponse])
async def list_dossiers(user: dict = Depends(get_current_user)):
    dossiers = await db.dossiers.find({"user_id": user["id"]}, {"_id": 0}).to_list(100)
    if not dossiers:
        return []
    # Single aggregation query instead of N+1 count_documents
    dossier_ids = [d["id"] for d in dossiers]
    piece_counts: dict = {}
    pipeline = [
        {"$match": {"dossier_id": {"$in": dossier_ids}}},
        {"$group": {"_id": "$dossier_id", "count": {"$sum": 1}}},
    ]
    async for row in db.pieces.aggregate(pipeline):
        piece_counts[row["_id"]] = row["count"]
    return [
        DossierResponse(**d, piece_count=piece_counts.get(d["id"], 0))
        for d in dossiers
    ]

@api_router.get("/dossiers/{dossier_id}", response_model=DossierResponse)
async def get_dossier(dossier_id: str, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    count = await db.pieces.count_documents({"dossier_id": dossier_id})
    return DossierResponse(**dossier, piece_count=count)

@api_router.put("/dossiers/{dossier_id}", response_model=DossierResponse)
async def update_dossier(dossier_id: str, data: DossierUpdate, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    update_data = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if data.title is not None:
        update_data["title"] = data.title
    if data.description is not None:
        update_data["description"] = data.description
    
    await db.dossiers.update_one({"id": dossier_id}, {"$set": update_data})
    dossier = await db.dossiers.find_one({"id": dossier_id}, {"_id": 0})
    count = await db.pieces.count_documents({"dossier_id": dossier_id})
    return DossierResponse(**dossier, piece_count=count)

@api_router.delete("/dossiers/{dossier_id}")
async def delete_dossier(dossier_id: str, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id}).to_list(1000)
    for piece in pieces:
        await storage.delete_file(piece["filename"])
    
    await db.pieces.delete_many({"dossier_id": dossier_id})
    await db.share_links.delete_many({"dossier_id": dossier_id})
    await db.dossiers.delete_one({"id": dossier_id})
    return {"message": "Dossier deleted"}

# ===================== FILE PROCESSING =====================

def convert_heic_to_jpg(filepath: Path) -> Path:
    """Convert HEIC to JPG"""
    try:
        from PIL import Image
        from pillow_heif import register_heif_opener
        register_heif_opener()
        
        img = Image.open(filepath)
        new_path = filepath.with_suffix('.jpg')
        img.convert('RGB').save(new_path, 'JPEG', quality=95)
        filepath.unlink()
        return new_path
    except Exception as e:
        logger.error(f"HEIC conversion error: {e}")
        return filepath

def extract_text_from_docx(filepath: Path) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_doc(filepath: Path) -> str:
    """Extract text from DOC file using antiword or fallback"""
    try:
        result = subprocess.run(['antiword', str(filepath)], capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            return result.stdout
    except Exception:
        pass
    return ""

def extract_text_from_pdf(filepath: Path) -> str:
    """Extract text from PDF"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages[:50]:
            text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""

# ===================== PIECE ROUTES =====================

@api_router.post("/dossiers/{dossier_id}/check-duplicate")
async def check_duplicate(
    dossier_id: str, 
    file: UploadFile = File(...), 
    user: dict = Depends(get_current_user)
):
    """Check if a file is a duplicate before uploading"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    content = await file.read()
    file_hash = compute_file_hash(content)
    
    existing = await db.pieces.find_one(
        {"dossier_id": dossier_id, "file_hash": file_hash},
        {"_id": 0}
    )
    
    if existing:
        return DuplicateCheckResponse(
            is_duplicate=True,
            existing_piece_id=existing["id"],
            existing_piece_numero=existing["numero"]
        )
    
    return DuplicateCheckResponse(is_duplicate=False)

@api_router.post("/dossiers/{dossier_id}/pieces", response_model=PieceResponse)
async def upload_piece(
    dossier_id: str, 
    file: UploadFile = File(...),
    force_upload: bool = Query(False, description="Upload even if duplicate"),
    source: str = Query(None, description="Source of file: 'camera' for photos taken with camera"),
    user: dict = Depends(get_current_user)
):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Check plan limits for total pieces
    dossier_ids = [d["id"] async for d in db.dossiers.find({"user_id": user["id"]}, {"id": 1})]
    total_pieces = await db.pieces.count_documents({"dossier_id": {"$in": dossier_ids}})
    await check_plan_limit(user, "max_total_pieces", total_pieces)
    
    # Check plan limits for pieces in this dossier
    dossier_pieces = await db.pieces.count_documents({"dossier_id": dossier_id})
    await check_plan_limit(user, "max_pieces_per_dossier", dossier_pieces)
    
    # Read file content ONCE and store in memory
    content = await file.read()
    file_size = len(content)
    
    # Reject empty files
    if file_size == 0:
        raise HTTPException(status_code=400, detail="Le fichier est vide (0 bytes)")
    
    if file_size > config.MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=413, detail=f"Fichier trop volumineux (max {config.MAX_FILE_SIZE_MB} MB)")
    
    # Compute hash for duplicate detection on raw bytes
    file_hash = compute_file_hash(content)
    content_type = file.content_type or "application/octet-stream"
    
    # Check for duplicate using hash AND size
    existing = await db.pieces.find_one({
        "dossier_id": dossier_id, 
        "file_hash": file_hash,
        "file_size": file_size
    })
    
    is_duplicate = False
    if existing:
        is_duplicate = True
        if not force_upload:
            # Return detailed error for frontend to display
            raise HTTPException(
                status_code=409, 
                detail={
                    "message": f"Fichier identique déjà présent (Pièce {existing['numero']})",
                    "existing_piece_id": existing["id"],
                    "existing_piece_numero": existing["numero"],
                    "existing_filename": existing["original_filename"]
                }
            )
    
    # Get next piece number
    last_piece = await db.pieces.find_one({"dossier_id": dossier_id}, sort=[("numero", -1)])
    next_numero = (last_piece["numero"] + 1) if last_piece else 1
    
    # Determine file type
    ext = Path(file.filename).suffix.lower()
    file_type_map = {
        ".pdf": "pdf", ".jpg": "image", ".jpeg": "image", ".png": "image",
        ".docx": "docx", ".doc": "doc", ".heic": "heic", ".heif": "heic",
        ".txt": "text"
    }
    file_type = file_type_map.get(ext, "other")
    
    # Save file - use stored content bytes
    piece_id = str(uuid.uuid4())
    filename = f"{piece_id}{ext}"
    
    # Convert HEIC if needed (in memory before saving)
    if file_type == "heic":
        try:
            from PIL import Image
            from pillow_heif import register_heif_opener
            register_heif_opener()
            img = Image.open(io.BytesIO(content))
            jpg_buffer = io.BytesIO()
            img.convert('RGB').save(jpg_buffer, 'JPEG', quality=95)
            content = jpg_buffer.getvalue()
            filename = f"{piece_id}.jpg"
            file_size = len(content)
            file_type = "image"
        except Exception as e:
            logger.error(f"HEIC conversion error: {e}")
    
    # Save to storage backend (GridFS/S3/local)
    await storage.save_file(content, filename)
    
    now = datetime.now(timezone.utc).isoformat()
    piece_doc = {
        "id": piece_id,
        "dossier_id": dossier_id,
        "numero": next_numero,
        "filename": filename,
        "original_filename": file.filename,
        "file_type": file_type,
        "file_size": file_size,
        "content_type": content_type,
        "file_hash": file_hash,
        "is_duplicate": is_duplicate,
        "source": source if source else "upload",  # 'camera' or 'upload'
        "status": "a_verifier",
        "analysis_status": "pending",
        "analysis_error": None,
        "analysis_queued_at": None,
        "analysis_started_at": None,
        "analysis_completed_at": None,
        "ai_proposal": None,
        "validated_data": None,
        "extracted_text": None,
        "created_at": now,
        "updated_at": now
    }
    await db.pieces.insert_one(piece_doc)
    
    logger.info(f"Uploaded piece {piece_id}: {file.filename}, size={file_size}, type={file_type}")
    
    return PieceResponse(**{k: v for k, v in piece_doc.items() if k not in ["extracted_text", "content_type"]})

@api_router.get("/dossiers/{dossier_id}/pieces", response_model=List[PieceResponse])
async def list_pieces(
    dossier_id: str, 
    filter_duplicates: bool = Query(False),
    filter_errors: bool = Query(False),
    user: dict = Depends(get_current_user)
):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    query = {"dossier_id": dossier_id}
    if filter_duplicates:
        query["is_duplicate"] = True
    if filter_errors:
        query["analysis_status"] = "error"
    
    pieces = await db.pieces.find(query, {"_id": 0, "extracted_text": 0}).sort("numero", 1).to_list(1000)
    
    # Check file existence in storage for each piece
    result = []
    for p in pieces:
        missing = not await storage.file_exists(p["filename"])
        p["file_missing"] = missing
        result.append(PieceResponse(**p))
    return result

@api_router.get("/pieces/{piece_id}", response_model=PieceResponse)
async def get_piece(piece_id: str, user: dict = Depends(get_current_user)):
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0, "extracted_text": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    piece["file_missing"] = not await storage.file_exists(piece["filename"])
    return PieceResponse(**piece)

@api_router.get("/pieces/{piece_id}/file")
async def get_piece_file(piece_id: str, user: dict = Depends(get_current_user)):
    """Download piece file with proper authentication"""
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    try:
        content = await storage.get_file(piece["filename"])
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="File not found")
    
    return Response(
        content=content,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{piece["original_filename"]}"'}
    )

@api_router.get("/pieces/{piece_id}/preview")
async def preview_piece_file(piece_id: str, user: dict = Depends(get_current_user)):
    """Get file content for inline preview (PDF/images)"""
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    try:
        content = await storage.get_file(piece["filename"])
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="File not found")
    
    # Determine content type
    ext = Path(piece["filename"]).suffix.lower()
    content_types = {
        ".pdf": "application/pdf",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
    }
    content_type = content_types.get(ext, "application/octet-stream")
    
    return Response(
        content=content,
        media_type=content_type,
        headers={
            "Content-Disposition": f'inline; filename="{piece["original_filename"]}"'
        }
    )

# ===================== ANALYSIS QUEUE =====================

@api_router.post("/dossiers/{dossier_id}/queue-analysis")
async def queue_analysis(
    dossier_id: str, 
    request: QueueAnalysisRequest,
    user: dict = Depends(get_current_user)
):
    """Queue pieces for analysis (instead of analyzing all at once)"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Build query
    query = {"dossier_id": dossier_id}
    if request.piece_ids:
        query["id"] = {"$in": request.piece_ids}
    else:
        # Only queue pieces that haven't been analyzed or had errors
        query["analysis_status"] = {"$in": ["pending", "error"]}
    
    now = datetime.now(timezone.utc).isoformat()
    
    result = await db.pieces.update_many(
        query,
        {"$set": {
            "analysis_status": "queued",
            "analysis_queued_at": now,
            "analysis_error": None,
            "updated_at": now
        }}
    )
    
    return {
        "message": f"{result.modified_count} pièces ajoutées à la file d'attente",
        "queued_count": result.modified_count
    }

@api_router.post("/dossiers/{dossier_id}/queue-failed")
async def queue_failed_analyses(dossier_id: str, user: dict = Depends(get_current_user)):
    """Re-queue all failed analyses"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    now = datetime.now(timezone.utc).isoformat()
    
    result = await db.pieces.update_many(
        {"dossier_id": dossier_id, "analysis_status": "error"},
        {"$set": {
            "analysis_status": "queued",
            "analysis_queued_at": now,
            "analysis_error": None,
            "updated_at": now
        }}
    )
    
    return {
        "message": f"{result.modified_count} pièces en échec remises en file d'attente",
        "queued_count": result.modified_count
    }

@api_router.post("/dossiers/{dossier_id}/process-queue")
async def process_analysis_queue(dossier_id: str, user: dict = Depends(get_current_user)):
    """Process queued analyses with rate limiting"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Check rate limit (simple implementation)
    lock_key = f"analysis_{dossier_id}"
    now = datetime.now(timezone.utc)
    
    if lock_key in analysis_locks:
        last_run = analysis_locks[lock_key]
        if (now - last_run).total_seconds() < ANALYSIS_RATE_LIMIT_SECONDS:
            return {"message": "Veuillez patienter quelques secondes avant de relancer", "processed": 0}
    
    analysis_locks[lock_key] = now
    
    # Count currently analyzing
    analyzing_count = await db.pieces.count_documents({
        "dossier_id": dossier_id,
        "analysis_status": "analyzing"
    })
    
    # Calculate how many we can process
    slots_available = MAX_CONCURRENT_ANALYSES - analyzing_count
    if slots_available <= 0:
        return {"message": "Analyses en cours, veuillez patienter", "processed": 0}
    
    # Get queued pieces
    queued_pieces = await db.pieces.find(
        {"dossier_id": dossier_id, "analysis_status": "queued"},
        {"_id": 0}
    ).sort("analysis_queued_at", 1).limit(slots_available).to_list(slots_available)
    
    processed = 0
    for piece in queued_pieces:
        try:
            # Mark as analyzing
            await db.pieces.update_one(
                {"id": piece["id"]},
                {"$set": {
                    "analysis_status": "analyzing",
                    "analysis_started_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            
            # Perform analysis
            filepath = await storage.get_temp_filepath(piece["filename"])
            try:
                ai_proposal = await analyze_document_with_ai(
                    filepath, piece["file_type"], piece["original_filename"], piece["id"]
                )
            finally:
                # Clean up temp file (only if it's a real temp file, not local storage)
                if not isinstance(storage, LocalStorage) and filepath.exists():
                    filepath.unlink()
            
            # Update with results
            await db.pieces.update_one(
                {"id": piece["id"]},
                {"$set": {
                    "ai_proposal": ai_proposal,
                    "analysis_status": "complete",
                    "analysis_completed_at": datetime.now(timezone.utc).isoformat(),
                    "analysis_error": None,
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )
            processed += 1
            
        except Exception as e:
            logger.error(f"Analysis failed for piece {piece['id']}: {e}")
            await db.pieces.update_one(
                {"id": piece["id"]},
                {"$set": {
                    "analysis_status": "error",
                    "analysis_error": str(e)[:500],
                    "analysis_completed_at": datetime.now(timezone.utc).isoformat(),
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )
    
    # Check if more in queue
    remaining = await db.pieces.count_documents({
        "dossier_id": dossier_id,
        "analysis_status": "queued"
    })
    
    return {
        "message": f"{processed} analyse(s) terminée(s)",
        "processed": processed,
        "remaining_in_queue": remaining
    }

@api_router.get("/dossiers/{dossier_id}/queue-status")
async def get_queue_status(dossier_id: str, user: dict = Depends(get_current_user)):
    """Get current analysis queue status"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pending = await db.pieces.count_documents({"dossier_id": dossier_id, "analysis_status": "pending"})
    queued = await db.pieces.count_documents({"dossier_id": dossier_id, "analysis_status": "queued"})
    analyzing = await db.pieces.count_documents({"dossier_id": dossier_id, "analysis_status": "analyzing"})
    complete = await db.pieces.count_documents({"dossier_id": dossier_id, "analysis_status": "complete"})
    error = await db.pieces.count_documents({"dossier_id": dossier_id, "analysis_status": "error"})
    
    return {
        "pending": pending,
        "queued": queued,
        "analyzing": analyzing,
        "complete": complete,
        "error": error,
        "total": pending + queued + analyzing + complete + error
    }

@api_router.post("/pieces/{piece_id}/analyze", response_model=PieceResponse)
async def analyze_piece(piece_id: str, user: dict = Depends(get_current_user)):
    """Analyze a single piece (with rate limiting)"""
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Rate limit check
    lock_key = f"piece_{piece_id}"
    now = datetime.now(timezone.utc)
    
    if lock_key in analysis_locks:
        last_run = analysis_locks[lock_key]
        if (now - last_run).total_seconds() < ANALYSIS_RATE_LIMIT_SECONDS:
            raise HTTPException(status_code=429, detail="Veuillez patienter avant de relancer l'analyse")
    
    analysis_locks[lock_key] = now
    
    # Set status to analyzing
    await db.pieces.update_one(
        {"id": piece_id}, 
        {"$set": {
            "analysis_status": "analyzing",
            "analysis_started_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    try:
        filepath = await storage.get_temp_filepath(piece["filename"])
        try:
            ai_proposal = await analyze_document_with_ai(filepath, piece["file_type"], piece["original_filename"], piece_id)
        finally:
            if not isinstance(storage, LocalStorage) and filepath.exists():
                filepath.unlink()
        
        await db.pieces.update_one(
            {"id": piece_id},
            {"$set": {
                "ai_proposal": ai_proposal,
                "analysis_status": "complete",
                "analysis_completed_at": datetime.now(timezone.utc).isoformat(),
                "analysis_error": None,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
    except Exception as e:
        logger.error(f"Analysis error: {e}")
        await db.pieces.update_one(
            {"id": piece_id}, 
            {"$set": {
                "analysis_status": "error",
                "analysis_error": str(e)[:500],
                "analysis_completed_at": datetime.now(timezone.utc).isoformat()
            }}
        )
    finally:
        # Clean up lock to prevent memory leak
        analysis_locks.pop(lock_key, None)
    
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0, "extracted_text": 0})
    return PieceResponse(**piece)

@api_router.post("/pieces/{piece_id}/reanalyze", response_model=PieceResponse)
async def reanalyze_piece(piece_id: str, user: dict = Depends(get_current_user)):
    """Re-analyze a piece (clear previous and retry)"""
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Rate limit check
    lock_key = f"piece_{piece_id}"
    now = datetime.now(timezone.utc)
    
    if lock_key in analysis_locks:
        last_run = analysis_locks[lock_key]
        if (now - last_run).total_seconds() < ANALYSIS_RATE_LIMIT_SECONDS:
            raise HTTPException(status_code=429, detail="Veuillez patienter avant de relancer l'analyse")
    
    analysis_locks[lock_key] = now
    
    # Clear previous analysis
    await db.pieces.update_one(
        {"id": piece_id},
        {"$set": {
            "ai_proposal": None, 
            "analysis_status": "analyzing", 
            "extracted_text": None,
            "analysis_error": None,
            "analysis_started_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    try:
        filepath = await storage.get_temp_filepath(piece["filename"])
        try:
            ai_proposal = await analyze_document_with_ai(filepath, piece["file_type"], piece["original_filename"], piece_id)
        finally:
            if not isinstance(storage, LocalStorage) and filepath.exists():
                filepath.unlink()
        
        await db.pieces.update_one(
            {"id": piece_id},
            {"$set": {
                "ai_proposal": ai_proposal,
                "analysis_status": "complete",
                "analysis_completed_at": datetime.now(timezone.utc).isoformat(),
                "analysis_error": None,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
    except Exception as e:
        logger.error(f"Re-analysis error: {e}")
        await db.pieces.update_one(
            {"id": piece_id}, 
            {"$set": {
                "analysis_status": "error",
                "analysis_error": str(e)[:500],
                "analysis_completed_at": datetime.now(timezone.utc).isoformat()
            }}
        )
    finally:
        # Clean up lock to prevent memory leak
        analysis_locks.pop(lock_key, None)
    
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0, "extracted_text": 0})
    return PieceResponse(**piece)

@api_router.post("/pieces/{piece_id}/validate", response_model=PieceResponse)
async def validate_piece(piece_id: str, data: PieceValidation, user: dict = Depends(get_current_user)):
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    await db.pieces.update_one(
        {"id": piece_id},
        {"$set": {
            "validated_data": data.model_dump(),
            "status": "pret",
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0, "extracted_text": 0})
    return PieceResponse(**piece)

@api_router.post("/pieces/{piece_id}/reupload", response_model=PieceResponse)
async def reupload_piece_file(
    piece_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user)
):
    """Re-upload a file for an existing piece (when file is missing)"""
    piece = await db.pieces.find_one({"id": piece_id}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    content = await file.read()
    file_size = len(content)
    
    if file_size > config.MAX_FILE_SIZE_MB * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"Fichier trop volumineux (max {config.MAX_FILE_SIZE_MB} Mo)")
    
    # Delete old file if it exists
    await storage.delete_file(piece["filename"])
    
    # Save new file with the same filename key
    ext = Path(file.filename).suffix.lower() if file.filename else Path(piece["original_filename"]).suffix.lower()
    new_filename = f"{piece_id}{ext}"
    
    # Handle HEIC
    if ext in ('.heic', '.heif'):
        try:
            from PIL import Image
            from pillow_heif import register_heif_opener
            register_heif_opener()
            img = Image.open(io.BytesIO(content))
            jpg_buffer = io.BytesIO()
            img.convert('RGB').save(jpg_buffer, 'JPEG', quality=95)
            content = jpg_buffer.getvalue()
            new_filename = f"{piece_id}.jpg"
            file_size = len(content)
        except Exception as e:
            logger.error(f"HEIC conversion error: {e}")
    
    await storage.save_file(content, new_filename)
    
    file_hash = compute_file_hash(content)
    
    await db.pieces.update_one(
        {"id": piece_id},
        {"$set": {
            "filename": new_filename,
            "original_filename": file.filename or piece["original_filename"],
            "file_size": file_size,
            "file_hash": file_hash,
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    updated = await db.pieces.find_one({"id": piece_id}, {"_id": 0, "extracted_text": 0})
    updated["file_missing"] = False
    return PieceResponse(**updated)

@api_router.post("/dossiers/{dossier_id}/bulk-reupload")
async def bulk_reupload_pieces(
    dossier_id: str,
    files: List[UploadFile] = File(...),
    user: dict = Depends(get_current_user)
):
    """Bulk re-upload: drop multiple files, auto-match to missing pieces.
    Matching strategy (in order):
    1. Exact original filename match
    2. Content hash match (same file, different name)
    3. Stem match (filename without extension)
    """
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Get all pieces with missing files in this dossier
    all_pieces = await db.pieces.find({"dossier_id": dossier_id}, {"_id": 0}).to_list(1000)
    missing_pieces = []
    for p in all_pieces:
        if not await storage.file_exists(p["filename"]):
            missing_pieces.append(p)
    
    if not missing_pieces:
        return {"matched": 0, "unmatched_files": [f.filename for f in files], "message": "Aucune pièce manquante dans ce dossier"}
    
    # Build lookups
    by_name = {}      # original_filename (lowercase) -> piece
    by_hash = {}      # file_hash -> piece
    by_stem = {}      # filename stem (no ext, lowercase) -> piece
    remaining = {}    # piece_id -> piece (for tracking unmatched)
    
    for p in missing_pieces:
        pid = p["id"]
        remaining[pid] = p
        name_key = p["original_filename"].strip().lower()
        by_name[name_key] = p
        stem_key = Path(name_key).stem
        if stem_key not in by_stem:
            by_stem[stem_key] = p
        if p.get("file_hash"):
            by_hash[p["file_hash"]] = p
    
    # Read all uploaded files first (we need content for hash matching)
    uploaded = []
    for upload_file in files:
        content = await upload_file.read()
        uploaded.append({
            "filename": upload_file.filename or "",
            "content": content,
            "hash": compute_file_hash(content),
            "size": len(content)
        })
    
    matched = 0
    unmatched_files = []
    restored_pieces = []
    
    async def save_match(piece, ufile, match_method):
        nonlocal matched
        content = ufile["content"]
        file_size = ufile["size"]
        ext = Path(ufile["filename"]).suffix.lower() if ufile["filename"] else Path(piece["original_filename"]).suffix.lower()
        new_filename = f"{piece['id']}{ext}"
        
        # Handle HEIC
        if ext in ('.heic', '.heif'):
            try:
                from PIL import Image
                from pillow_heif import register_heif_opener
                register_heif_opener()
                img = Image.open(io.BytesIO(content))
                jpg_buffer = io.BytesIO()
                img.convert('RGB').save(jpg_buffer, 'JPEG', quality=95)
                content = jpg_buffer.getvalue()
                new_filename = f"{piece['id']}.jpg"
                file_size = len(content)
            except Exception as e:
                logger.error(f"HEIC conversion error: {e}")
        
        await storage.save_file(content, new_filename)
        file_hash = compute_file_hash(content)
        
        await db.pieces.update_one(
            {"id": piece["id"]},
            {"$set": {
                "filename": new_filename,
                "file_size": file_size,
                "file_hash": file_hash,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        
        matched += 1
        restored_pieces.append({
            "piece_id": piece["id"],
            "original_filename": piece["original_filename"],
            "numero": piece["numero"],
            "matched_with": ufile["filename"],
            "match_method": match_method
        })
        # Remove from all lookups
        remaining.pop(piece["id"], None)
        name_key = piece["original_filename"].strip().lower()
        by_name.pop(name_key, None)
        stem_key = Path(name_key).stem
        by_stem.pop(stem_key, None)
        if piece.get("file_hash"):
            by_hash.pop(piece["file_hash"], None)
    
    # Pass 1: Exact filename match
    still_to_match = []
    for ufile in uploaded:
        fname = ufile["filename"].strip().lower()
        piece = by_name.get(fname)
        if piece and piece["id"] in remaining:
            await save_match(piece, ufile, "nom_exact")
        else:
            still_to_match.append(ufile)
    
    # Pass 2: Content hash match (handles renamed files)
    still_to_match2 = []
    for ufile in still_to_match:
        piece = by_hash.get(ufile["hash"])
        if piece and piece["id"] in remaining:
            await save_match(piece, ufile, "contenu_identique")
        else:
            still_to_match2.append(ufile)
    
    # Pass 3: Stem match (filename without extension)
    for ufile in still_to_match2:
        fname = ufile["filename"].strip().lower()
        stem = Path(fname).stem
        piece = by_stem.get(stem)
        if piece and piece["id"] in remaining:
            await save_match(piece, ufile, "nom_approchant")
        else:
            unmatched_files.append(ufile["filename"])
    
    still_missing = [{"piece_id": p["id"], "original_filename": p["original_filename"], "numero": p["numero"]} for p in remaining.values()]
    
    return {
        "matched": matched,
        "restored": restored_pieces,
        "unmatched_files": unmatched_files,
        "still_missing": still_missing,
        "total_missing_before": len(missing_pieces),
        "total_missing_after": len(still_missing)
    }

@api_router.get("/dossiers/{dossier_id}/missing-pieces")
async def get_missing_pieces(dossier_id: str, user: dict = Depends(get_current_user)):
    """Get list of pieces with missing files in a dossier"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id}, {"_id": 0, "extracted_text": 0}).to_list(1000)
    missing = []
    for p in pieces:
        if not await storage.file_exists(p["filename"]):
            missing.append({
                "piece_id": p["id"],
                "numero": p["numero"],
                "original_filename": p["original_filename"],
                "file_type": p["file_type"],
                "validated_titre": p.get("validated_data", {}).get("titre", "") if p.get("validated_data") else "",
                "ai_titre": p.get("ai_proposal", {}).get("titre", "") if p.get("ai_proposal") else ""
            })
    return {"missing_count": len(missing), "missing_pieces": missing}

@api_router.delete("/pieces/{piece_id}")
async def delete_piece(piece_id: str, user: dict = Depends(get_current_user)):
    piece = await db.pieces.find_one({"id": piece_id})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    dossier = await db.dossiers.find_one({"id": piece["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Access denied")
    
    await storage.delete_file(piece["filename"])
    
    await db.pieces.delete_one({"id": piece_id})
    return {"message": "Piece deleted"}

@api_router.post("/dossiers/{dossier_id}/pieces/delete-many")
async def delete_many_pieces(
    dossier_id: str, 
    request: DeleteManyRequest,
    user: dict = Depends(get_current_user)
):
    """Delete multiple pieces at once"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Get pieces that belong to this dossier
    pieces = await db.pieces.find({
        "id": {"$in": request.piece_ids},
        "dossier_id": dossier_id
    }).to_list(len(request.piece_ids))
    
    deleted_count = 0
    for piece in pieces:
        await storage.delete_file(piece["filename"])
        await db.pieces.delete_one({"id": piece["id"]})
        deleted_count += 1
    
    return {"message": f"{deleted_count} pièces supprimées", "deleted_count": deleted_count}

@api_router.post("/dossiers/{dossier_id}/pieces/delete-errors")
async def delete_error_pieces(dossier_id: str, user: dict = Depends(get_current_user)):
    """Delete all pieces with analysis errors"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({
        "dossier_id": dossier_id,
        "analysis_status": "error"
    }).to_list(1000)
    
    deleted_count = 0
    for piece in pieces:
        await storage.delete_file(piece["filename"])
        await db.pieces.delete_one({"id": piece["id"]})
        deleted_count += 1
    
    return {"message": f"{deleted_count} pièces en erreur supprimées", "deleted_count": deleted_count}

@api_router.post("/dossiers/{dossier_id}/renumber")
async def renumber_pieces(dossier_id: str, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id}).sort("numero", 1).to_list(1000)
    for i, piece in enumerate(pieces, 1):
        await db.pieces.update_one({"id": piece["id"]}, {"$set": {"numero": i}})
    
    return {"message": "Pieces renumbered"}

# ===================== AI ANALYSIS =====================

async def analyze_document_with_ai(filepath: Path, file_type: str, original_filename: str, piece_id: str) -> dict:
    """Analyze document using Gemini Vision"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
    
    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        logger.error("EMERGENT_LLM_KEY not set")
        raise Exception("Configuration serveur incomplète (clé API manquante)")
    
    try:
        # Extract text for DOCX/DOC files first
        extracted_text = None
        if file_type == "docx":
            extracted_text = extract_text_from_docx(filepath)
        elif file_type == "doc":
            extracted_text = extract_text_from_doc(filepath)
        elif file_type == "pdf":
            extracted_text = extract_text_from_pdf(filepath)
        
        # Store extracted text
        if extracted_text:
            await db.pieces.update_one({"id": piece_id}, {"$set": {"extracted_text": extracted_text[:50000]}})
        
        system_message = """Tu es un assistant juridique expert. Analyse ce document juridique et extrais les informations suivantes de manière structurée.

IMPORTANT:
- N'invente JAMAIS d'information. Si tu ne trouves pas une information, indique "Non identifié".
- Pour chaque information, indique un niveau de confiance: "faible", "moyen", ou "fort".
- Cite un extrait du document (max 200 caractères) qui justifie ta proposition.
- Classifie le document de manière NEUTRE selon 3 axes :
  • tags_thematiques : zéro ou plusieurs parmi cette liste fermée :
    "famille", "travail", "sante", "finances", "logement", "violence",
    "harcelement", "administratif", "scolaire", "succession", "consommation"
  • sujets_concernes : qui est concerné, parmi cette liste fermée :
    "utilisateur", "conjoint", "ex-conjoint", "enfant", "employeur",
    "bailleur", "voisin", "administration", "medecin", "police", "tiers"
  • nature_document : un seul mot parmi :
    "officiel" (acte/décision d'autorité), "prive" (courrier/SMS/email perso),
    "temoignage" (attestation, déclaration tiers), "medical" (certif/ordonnance),
    "financier" (facture/relevé/contrat), "autre"
- Tu DOIS rester strictement factuel et neutre. Ne juge pas, ne qualifie pas, n'interprete pas. Ne signale pas l'utilisateur comme victime ou coupable.

Réponds UNIQUEMENT en JSON valide avec cette structure exacte:
{
  "type_piece": "plainte|main_courante|certificat_medical|attestation|sms|conclusions|assignation|recit|facture|contrat|jugement|ordonnance|autre",
  "type_confidence": "faible|moyen|fort",
  "date_document": "YYYY-MM-DD ou null si non trouvée",
  "date_confidence": "faible|moyen|fort",
  "titre": "titre clair et standardisé",
  "titre_confidence": "faible|moyen|fort",
  "resume_qui": "personnes impliquées",
  "resume_quoi": "fait ou motif principal",
  "resume_ou": "lieu si mentionné ou null",
  "resume_element_cle": "diagnostic, menace, refus, constat, montant, décision, etc.",
  "mots_cles": ["mot1", "mot2", "mot3"],
  "extrait_justificatif": "extrait du document justifiant l'analyse (max 200 car.)",
  "tags_thematiques": ["famille", "violence"],
  "sujets_concernes": ["utilisateur", "conjoint"],
  "nature_document": "officiel"
}"""

        chat = LlmChat(
            api_key=api_key,
            session_id=f"analysis-{uuid.uuid4()}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.5-flash")
        
        # For text-based files with extracted text, send the text
        if extracted_text and len(extracted_text) > 100:
            user_message = UserMessage(
                text=f"Analyse ce document juridique (nom: {original_filename}).\n\nContenu extrait:\n{extracted_text[:30000]}"
            )
        else:
            # For images and PDFs, use vision
            mime_map = {
                "pdf": "application/pdf",
                "image": "image/jpeg" if filepath.suffix.lower() in [".jpg", ".jpeg"] else "image/png",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            mime_type = mime_map.get(file_type, "application/octet-stream")
            
            file_content = FileContentWithMimeType(
                file_path=str(filepath),
                mime_type=mime_type
            )
            
            user_message = UserMessage(
                text=f"Analyse ce document juridique (nom original: {original_filename}). Extrais toutes les informations pertinentes.",
                file_contents=[file_content]
            )
        
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        response_text = response.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        proposal = json.loads(response_text.strip())
        # Normalise tags_thematiques vers les 5 domaines juridiques canoniques.
        try:
            from piece_classifier import normalize_themes
            proposal["tags_thematiques"] = normalize_themes(proposal.get("tags_thematiques") or [])
        except Exception:
            pass
        return proposal
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {e}")
        raise Exception(f"Erreur de parsing de la réponse IA")
    except Exception as e:
        logger.error(f"AI analysis error: {e}")
        raise

# ===================== CHRONOLOGY & EXPORTS =====================

# Suggestions d'aide à la constitution selon les 5 domaines juridiques canoniques.
THEME_HINTS = {
    "PÉNAL": [
        "Plaintes ou mains courantes déposées",
        "Certificats médicaux constatant des lésions",
        "Captures d'écran horodatées (SMS, emails, réseaux sociaux)",
        "Témoignages écrits de proches, voisins ou témoins directs",
    ],
    "CIVIL_FAMILLE": [
        "Acte de mariage / livret de famille",
        "Jugements ou ordonnances antérieurs (JAF, divorce, garde)",
        "Échanges écrits avec la partie adverse (SMS, emails, courriers)",
        "Acte de décès, testament et inventaire (succession)",
    ],
    "IMMOBILIER_LOGEMENT": [
        "Bail de location et annexes",
        "Quittances de loyer / appels de charges",
        "État des lieux d'entrée et de sortie",
        "Courriers recommandés avec accusé de réception",
    ],
    "TRAVAIL": [
        "Contrat de travail et avenants",
        "Bulletins de salaire des 12 derniers mois",
        "Échanges écrits avec l'employeur (emails, courriers)",
        "Lettre de licenciement / rupture conventionnelle",
    ],
    "ADMINISTRATIF": [
        "Décisions administratives (notifications, refus)",
        "Justificatifs de démarches effectuées",
        "Relevés bancaires et factures pour les volets financiers",
        "Bulletins ou rapports scolaires si volet enfants",
    ],
}


def _piece_classification(piece: dict) -> dict:
    """Récupère tags/sujets/nature/sous-domaine/source depuis validated_data en priorité, sinon ai_proposal."""
    v = piece.get("validated_data") or {}
    a = piece.get("ai_proposal") or {}
    return {
        "tags_thematiques": v.get("tags_thematiques") or a.get("tags_thematiques") or [],
        "sujets_concernes": v.get("sujets_concernes") or a.get("sujets_concernes") or [],
        "nature_document": v.get("nature_document") or a.get("nature_document"),
        "sous_domaine": v.get("sous_domaine") or a.get("sous_domaine"),
        "type_specifique": v.get("type_specifique") or a.get("type_specifique"),
        "source_qualifiee": v.get("source_qualifiee") or a.get("source_qualifiee"),
    }


@api_router.get("/dossiers/{dossier_id}/synthesis")
async def get_synthesis(dossier_id: str, user: dict = Depends(get_current_user)):
    """Synthèse factuelle des thèmes, sujets et natures détectés dans le dossier."""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    pieces = await db.pieces.find({"dossier_id": dossier_id}, {"_id": 0}).to_list(2000)
    return _compute_synthesis(pieces)


@api_router.post("/dossiers/{dossier_id}/reclassify")
async def reclassify_dossier_pieces(
    dossier_id: str,
    force: bool = False,
    user: dict = Depends(get_current_user),
):
    """
    Classification rétroactive (par mots-clés, SANS appel IA) des pièces existantes.
    Toutes les pièces sont traitées : on tente d'abord de normaliser les
    tags existants vers les 5 domaines juridiques canoniques ; si rien ne reste,
    on relance la détection par mots-clés (classify_piece).
    """
    from piece_classifier import normalize_themes, detect_subdomain, derive_source_qualifiee, _build_haystack  # noqa: PLC0415

    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")

    pieces = await db.pieces.find({"dossier_id": dossier_id}, {"_id": 0}).to_list(2000)
    updated = 0
    for p in pieces:
        v = p.get("validated_data") or {}
        a = p.get("ai_proposal") or {}
        existing = v.get("tags_thematiques") or a.get("tags_thematiques") or []
        normalized = normalize_themes(existing)

        if not normalized:
            # Fallback : détection par mots-clés (classify_piece applique aussi normalize_themes)
            cls = classify_piece(p)
            normalized = cls["tags_thematiques"]
            subjects = cls["sujets_concernes"]
            nature = cls["nature_document"]
            sous_domaine = cls["sous_domaine"]
            type_specifique = cls["type_specifique"]
            source_qualifiee = cls["source_qualifiee"]
        else:
            subjects = v.get("sujets_concernes") or a.get("sujets_concernes") or ["utilisateur"]
            nature = v.get("nature_document") or a.get("nature_document")
            if not nature:
                # Dérive depuis type_piece quand non stockée
                from piece_classifier import TYPE_TO_NATURE  # noqa: PLC0415
                type_piece = v.get("type_piece") or a.get("type_piece") or ""
                nature = TYPE_TO_NATURE.get(type_piece)
            # Toujours recalculer sous-domaine + type à partir du texte (texte stable)
            primary = normalized[0] if normalized else None
            sous_domaine, type_specifique = detect_subdomain(primary, _build_haystack(p))
            source_qualifiee = derive_source_qualifiee(nature)

        target_field = "validated_data" if v else "ai_proposal"
        target_doc = dict(v) if v else dict(a or {})
        target_doc["tags_thematiques"] = normalized
        target_doc["sujets_concernes"] = subjects
        target_doc["sous_domaine"] = sous_domaine
        target_doc["type_specifique"] = type_specifique
        target_doc["source_qualifiee"] = source_qualifiee
        if nature and not target_doc.get("nature_document"):
            target_doc["nature_document"] = nature

        await db.pieces.update_one(
            {"id": p["id"]},
            {"$set": {
                target_field: target_doc,
                "updated_at": datetime.now(timezone.utc).isoformat(),
            }},
        )
        updated += 1

    return {"ok": True, "updated": updated, "skipped": 0, "total": len(pieces)}


@api_router.get("/dossiers/{dossier_id}/chronology")
async def get_chronology(dossier_id: str, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id, "status": "pret"}, {"_id": 0}).to_list(1000)
    
    chronology = []
    for p in pieces:
        if p.get("validated_data"):
            date_doc = p["validated_data"].get("date_document")
            chronology.append({
                "piece_id": p["id"],
                "numero": p["numero"],
                "date": date_doc,
                "titre": p["validated_data"].get("titre", p["original_filename"]),
                "resume": {
                    "qui": p["validated_data"].get("resume_qui"),
                    "quoi": p["validated_data"].get("resume_quoi"),
                    "ou": p["validated_data"].get("resume_ou"),
                    "element_cle": p["validated_data"].get("resume_element_cle")
                },
                "type_piece": p["validated_data"].get("type_piece")
            })
    
    chronology.sort(key=lambda x: x["date"] or "9999-99-99")
    
    return {"dossier": dossier["title"], "dossier_id": dossier_id, "entries": chronology}

@api_router.get("/dossiers/{dossier_id}/export/csv")
async def export_csv(dossier_id: str, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id}, {"_id": 0}).sort("numero", 1).to_list(1000)
    
    csv_content = "N°;Type;Date;Titre;Résumé\n"
    for p in pieces:
        data = p.get("validated_data") or {}
        resume_parts = [data.get("resume_qui", ""), data.get("resume_quoi", ""), data.get("resume_element_cle", "")]
        resume = " - ".join([r for r in resume_parts if r])
        csv_content += f"Pièce {p['numero']};{data.get('type_piece', '')};{data.get('date_document', '')};{data.get('titre', p['original_filename'])};{resume}\n"
    
    return StreamingResponse(
        io.StringIO(csv_content),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=sommaire_{dossier_id}.csv"}
    )

@api_router.get("/dossiers/{dossier_id}/export/zip")
async def export_zip(dossier_id: str, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id}, {"_id": 0}).sort("numero", 1).to_list(1000)
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for p in pieces:
            try:
                file_content = await storage.get_file(p["filename"])
                ext = Path(p["original_filename"]).suffix
                arcname = f"Piece_{p['numero']}{ext}"
                zf.writestr(arcname, file_content)
            except FileNotFoundError:
                logger.warning(f"File not found for piece {p['id']}: {p['filename']}")
                continue
    
    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=dossier_{dossier_id}.zip"}
    )

# ===================== PDF EXPORT =====================

def format_date_fr(date_str: str) -> str:
    """Format date to DD/MM/YYYY"""
    if not date_str:
        return "Date non définie"
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return dt.strftime("%d/%m/%Y")
    except:
        try:
            parts = date_str.split("-")
            if len(parts) == 3:
                return f"{parts[2]}/{parts[1]}/{parts[0]}"
        except:
            pass
        return date_str

@api_router.get("/dossiers/{dossier_id}/export/pdf")
async def export_chronology_pdf(dossier_id: str, user: dict = Depends(get_current_user)):
    """Export chronology as professional PDF"""
    # Plan gating: PDF export reserved to paid plans (free plan blocked)
    plan = await get_user_plan(user)
    if not get_plan_limits(plan).can_export_pdf:
        raise HTTPException(
            status_code=403,
            detail={
                "error": "PLAN_LIMIT_EXCEEDED",
                "feature": "export_pdf",
                "message": "L'export PDF est réservé aux plans Essentiel et Sérénité. Passez au niveau supérieur pour télécharger votre dossier complet.",
                "upgrade_url": "/pricing"
            }
        )

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id, "status": "pret"}, {"_id": 0}).to_list(1000)
    
    # Build chronology
    chronology = []
    for p in pieces:
        if p.get("validated_data"):
            chronology.append({
                "numero": p["numero"],
                "date": p["validated_data"].get("date_document"),
                "titre": p["validated_data"].get("titre", p["original_filename"]),
                "type_piece": p["validated_data"].get("type_piece", "autre"),
                "resume": p["validated_data"]
            })
    chronology.sort(key=lambda x: x["date"] or "9999-99-99")
    
    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Title_Custom', fontName='Helvetica-Bold', fontSize=16, alignment=TA_CENTER, spaceAfter=20))
    styles.add(ParagraphStyle(name='Subtitle', fontName='Helvetica', fontSize=10, alignment=TA_CENTER, textColor=colors.grey, spaceAfter=30))
    styles.add(ParagraphStyle(name='Entry_Date', fontName='Helvetica-Bold', fontSize=11, textColor=colors.HexColor('#0369A1')))
    styles.add(ParagraphStyle(name='Entry_Title', fontName='Helvetica-Bold', fontSize=10, spaceAfter=5))
    styles.add(ParagraphStyle(name='Entry_Body', fontName='Helvetica', fontSize=9, alignment=TA_JUSTIFY, spaceAfter=3))
    styles.add(ParagraphStyle(name='Entry_Ref', fontName='Helvetica-Oblique', fontSize=8, textColor=colors.grey, spaceAfter=15))
    
    elements = []
    
    # Header
    elements.append(Paragraph(f"CHRONOLOGIE DES FAITS", styles['Title_Custom']))
    elements.append(Paragraph(f"Dossier : {dossier['title']}", styles['Subtitle']))
    elements.append(Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", styles['Subtitle']))
    elements.append(Spacer(1, 20))
    
    if not chronology:
        elements.append(Paragraph("Aucune pièce validée dans ce dossier.", styles['Entry_Body']))
    else:
        for entry in chronology:
            # Date
            date_str = format_date_fr(entry["date"])
            elements.append(Paragraph(f"• {date_str}", styles['Entry_Date']))
            
            # Title
            elements.append(Paragraph(entry["titre"], styles['Entry_Title']))
            
            # Resume
            resume = entry["resume"]
            resume_parts = []
            if resume.get("resume_qui"):
                resume_parts.append(f"<b>Qui :</b> {resume['resume_qui']}")
            if resume.get("resume_quoi"):
                resume_parts.append(f"<b>Quoi :</b> {resume['resume_quoi']}")
            if resume.get("resume_ou"):
                resume_parts.append(f"<b>Où :</b> {resume['resume_ou']}")
            if resume.get("resume_element_cle"):
                resume_parts.append(f"<b>Élément clé :</b> {resume['resume_element_cle']}")
            
            for part in resume_parts:
                elements.append(Paragraph(part, styles['Entry_Body']))
            
            # Reference
            elements.append(Paragraph(f"Référence : Pièce {entry['numero']}", styles['Entry_Ref']))
    
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=chronologie_{dossier_id}.pdf"}
    )

@api_router.get("/dossiers/{dossier_id}/export/docx")
async def export_chronology_docx(dossier_id: str, user: dict = Depends(get_current_user)):
    """Export chronology as DOCX narrative"""
    # Plan gating: DOCX export reserved to paid plans
    plan = await get_user_plan(user)
    if not get_plan_limits(plan).can_export_docx:
        raise HTTPException(
            status_code=403,
            detail={
                "error": "PLAN_LIMIT_EXCEEDED",
                "feature": "export_docx",
                "message": "L'export DOCX est réservé aux plans Essentiel et Sérénité. Passez au niveau supérieur pour télécharger votre dossier complet.",
                "upgrade_url": "/pricing"
            }
        )

    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": dossier_id, "status": "pret"}, {"_id": 0}).to_list(1000)
    
    # Build chronology
    chronology = []
    for p in pieces:
        if p.get("validated_data"):
            chronology.append({
                "numero": p["numero"],
                "date": p["validated_data"].get("date_document"),
                "titre": p["validated_data"].get("titre", p["original_filename"]),
                "type_piece": p["validated_data"].get("type_piece", "autre"),
                "resume": p["validated_data"]
            })
    chronology.sort(key=lambda x: x["date"] or "9999-99-99")
    
    # Create DOCX
    doc = Document()
    
    # Title
    title = doc.add_heading('CHRONOLOGIE NARRATIVE DES FAITS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Dossier : {dossier['title']}\nGénéré le {datetime.now().strftime('%d/%m/%Y')}")
    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    if not chronology:
        doc.add_paragraph("Aucune pièce validée dans ce dossier.")
    else:
        for entry in chronology:
            date_str = format_date_fr(entry["date"])
            resume = entry["resume"]
            
            # Build narrative paragraph
            para = doc.add_paragraph()
            
            # Date in bold
            run = para.add_run(f"Le {date_str}, ")
            run.bold = True
            
            # Build narrative text
            narrative_parts = []
            if resume.get("resume_qui"):
                narrative_parts.append(resume["resume_qui"])
            if resume.get("resume_quoi"):
                narrative_parts.append(resume["resume_quoi"].lower() if resume.get("resume_qui") else resume["resume_quoi"])
            if resume.get("resume_ou"):
                narrative_parts.append(f"à {resume['resume_ou']}")
            
            narrative = " ".join(narrative_parts) if narrative_parts else entry["titre"]
            para.add_run(narrative)
            
            if resume.get("resume_element_cle"):
                para.add_run(f". {resume['resume_element_cle']}")
            
            # Reference
            ref_run = para.add_run(f" (Pièce {entry['numero']})")
            ref_run.italic = True
            
            para.add_run(".")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=chronologie_narrative_{dossier_id}.docx"}
    )

# ===================== ASSISTANT DE RÉDACTION =====================

@api_router.post("/dossiers/{dossier_id}/assistant", response_model=AssistantResponse)
async def generate_document(dossier_id: str, request: AssistantRequest, user: dict = Depends(get_current_user)):
    """Generate document draft based on VALIDATED pieces only"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Check plan restrictions for assistant
    plan = await get_user_plan(user)
    limits = get_plan_limits(plan)

    # FREE plan: only expose_faits allowed, AND only 1 generation per dossier (lifetime)
    if plan == "free":
        if request.document_type != "expose_faits":
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "PLAN_LIMIT_EXCEEDED",
                    "feature": "assistant_document_type",
                    "message": "Le plan Découverte permet uniquement la génération d'un exposé des faits. Passez au plan Essentiel pour rédiger des courriers et bien plus.",
                    "plan": plan,
                    "allowed_types": ["expose_faits"],
                    "upgrade_url": "/pricing"
                }
            )
        # Already generated for THIS dossier?
        existing = await db.assistant_generations.count_documents({
            "user_id": user["id"],
            "dossier_id": dossier_id,
        })
        if existing >= 1:
            raise HTTPException(
                status_code=403,
                detail={
                    "error": "PLAN_LIMIT_EXCEEDED",
                    "feature": "assistant_per_dossier",
                    "message": "Le plan Découverte autorise un seul exposé des faits par dossier. Passez au plan Essentiel pour la rédaction illimitée.",
                    "plan": plan,
                    "upgrade_url": "/pricing"
                }
            )
    
    # Check daily usage limit
    assistant_uses = user.get("assistant_uses_today", 0)
    last_reset = user.get("assistant_last_reset")
    if last_reset:
        last_reset_date = datetime.fromisoformat(last_reset).date()
        if last_reset_date < datetime.now(timezone.utc).date():
            assistant_uses = 0
            # Reset the counter
            await db.users.update_one(
                {"id": user["id"]},
                {"$set": {"assistant_uses_today": 0, "assistant_last_reset": datetime.now(timezone.utc).isoformat()}}
            )
    
    await check_plan_limit(user, "assistant_daily_limit", assistant_uses)
    
    # Get validated pieces only
    query = {"dossier_id": dossier_id, "status": "pret"}
    if request.piece_ids:
        query["id"] = {"$in": request.piece_ids}
    
    pieces = await db.pieces.find(query, {"_id": 0}).sort("numero", 1).to_list(1000)
    
    # Filter by date if provided
    if request.date_start or request.date_end:
        filtered = []
        for p in pieces:
            date_doc = p.get("validated_data", {}).get("date_document")
            if date_doc:
                if request.date_start and date_doc < request.date_start:
                    continue
                if request.date_end and date_doc > request.date_end:
                    continue
            filtered.append(p)
        pieces = filtered
    
    if not pieces:
        return AssistantResponse(
            content="Aucune pièce validée disponible pour générer ce document.",
            pieces_used=[],
            warnings=["Aucune pièce sélectionnée ou toutes les pièces sont encore 'à vérifier'."]
        )
    
    # Build context from validated data ONLY
    context_parts = []
    pieces_used = []
    for p in pieces:
        vd = p.get("validated_data", {})
        if vd:
            pieces_used.append(p["numero"])
            entry = f"""
Pièce {p['numero']} - {vd.get('titre', p['original_filename'])}:
- Type: {vd.get('type_piece', 'non défini')}
- Date: {format_date_fr(vd.get('date_document'))}
- Qui: {vd.get('resume_qui', 'Non précisé')}
- Quoi: {vd.get('resume_quoi', 'Non précisé')}
- Où: {vd.get('resume_ou', 'Non précisé')}
- Élément clé: {vd.get('resume_element_cle', 'Non précisé')}
"""
            context_parts.append(entry)
    
    context = "\n".join(context_parts)
    
    # Labels for jurisdictions
    jurisdiction_labels = {
        "jaf": "Juge aux Affaires Familiales (JAF)",
        "penal": "Pénal",
        "prudhommes": "Prud'hommes",
        "administratif": "Administratif",
        "civil": "Civil (Tribunal judiciaire)",
        "commercial": "Commercial",
        "autre": "Juridiction à préciser"
    }
    
    jurisdiction_label = jurisdiction_labels.get(request.jurisdiction, "Juridiction à préciser") if request.jurisdiction else None
    
    # Document type prompts (agnostique du contentieux)
    prompts = {
        "expose_faits": f"""À partir des informations VALIDÉES suivantes, rédige un exposé des faits structuré et chronologique.

RÈGLES STRICTES:
- N'invente AUCUNE information
- Chaque fait doit citer sa source: (Pièce X)
- Si une information manque, écris "À confirmer"
- Style: juridique, factuel, neutre

Pièces validées:
{context}

Rédige l'exposé des faits:""",
        
        "chronologie_narrative": f"""À partir des informations VALIDÉES suivantes, rédige une chronologie narrative.

RÈGLES STRICTES:
- Un paragraphe par événement
- Format: "Le [date], [fait]. (Pièce X)"
- N'invente AUCUNE information
- Si une date est manquante, l'indiquer

Pièces validées:
{context}

Rédige la chronologie narrative:""",
        
        "courrier_avocat": f"""À partir des informations VALIDÉES suivantes, rédige un projet de courrier pour un avocat présentant la situation.

RÈGLES STRICTES:
- Chaque fait cité doit référencer sa source: (Pièce X)
- N'invente AUCUNE information
- Mentionner les points nécessitant clarification
- Style: professionnel, synthétique

Pièces validées:
{context}

Rédige le projet de courrier:""",
        
        "projet_requete": f"""À partir des informations VALIDÉES suivantes, rédige un projet de requête.

JURIDICTION CIBLÉE: {jurisdiction_label}

RÈGLES STRICTES:
- Commence par "PROJET DE REQUÊTE – {jurisdiction_label}"
- Chaque fait cité doit référencer sa source: (Pièce X)
- N'invente AUCUNE information
- Structure adaptée à la juridiction: Faits, Discussion, Par ces motifs (ou structure équivalente selon la juridiction)
- Les informations manquantes doivent être signalées avec "À confirmer"
- Reste factuel et neutre, les faits proviennent uniquement des pièces validées

Pièces validées:
{context}

Rédige le projet de requête pour la juridiction {jurisdiction_label}:"""
    }
    
    prompt = prompts.get(request.document_type, prompts["expose_faits"])
    
    api_key = os.environ.get("EMERGENT_LLM_KEY")
    if not api_key:
        return AssistantResponse(
            content="Service d'assistant non disponible (clé API manquante).",
            pieces_used=pieces_used,
            warnings=["Configuration serveur incomplète"]
        )
    
    try:
        chat = LlmChat(
            api_key=api_key,
            session_id=f"assistant-{uuid.uuid4()}",
            system_message="Tu es un assistant juridique. Tu rédiges des documents à partir d'informations VALIDÉES uniquement. Tu ne dois JAMAIS inventer d'information. Chaque fait doit citer sa source (Pièce X)."
        ).with_model("gemini", "gemini-2.5-flash")
        
        response = await chat.send_message(UserMessage(text=prompt))
        
        # Increment assistant usage counter
        await db.users.update_one(
            {"id": user["id"]},
            {"$inc": {"assistant_uses_today": 1}}
        )

        # Track generation per dossier (used for free-plan "1 per dossier" check)
        await db.assistant_generations.insert_one({
            "id": str(uuid.uuid4()),
            "user_id": user["id"],
            "dossier_id": dossier_id,
            "document_type": request.document_type,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        
        warnings = []
        if "À confirmer" in response:
            warnings.append("Certaines informations nécessitent confirmation (marquées 'À confirmer').")
        
        return AssistantResponse(
            content=response,
            pieces_used=pieces_used,
            warnings=warnings
        )
        
    except Exception as e:
        logger.error(f"Assistant error: {e}")
        return AssistantResponse(
            content="Erreur lors de la génération du document.",
            pieces_used=pieces_used,
            warnings=[str(e)]
        )

# ===================== SHARE LINKS =====================

@api_router.post("/dossiers/{dossier_id}/share", response_model=ShareLinkResponse)
async def create_share_link(dossier_id: str, data: ShareLinkCreate, user: dict = Depends(get_current_user)):
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    link_id = str(uuid.uuid4())
    token = str(uuid.uuid4())
    now = datetime.now(timezone.utc)
    expires_at = now + timedelta(days=data.expires_in_days)
    
    link_doc = {
        "id": link_id,
        "dossier_id": dossier_id,
        "token": token,
        "expires_at": expires_at.isoformat(),
        "created_at": now.isoformat(),
        "piece_ids": data.piece_ids  # None = all, [] = none, [...] = specific
    }
    await db.share_links.insert_one(link_doc)
    
    return ShareLinkResponse(**link_doc)

@api_router.get("/shared/{token}")
async def get_shared_dossier(token: str):
    """Public endpoint for lawyers"""
    link = await db.share_links.find_one({"token": token}, {"_id": 0})
    if not link:
        raise HTTPException(status_code=404, detail="Lien non trouvé")
    
    # Check if revoked
    if link.get("revoked"):
        raise HTTPException(status_code=410, detail="Ce lien a été révoqué par son propriétaire")
    
    # Check expiration
    expires = datetime.fromisoformat(link["expires_at"])
    if datetime.now(timezone.utc) > expires:
        raise HTTPException(status_code=410, detail="Ce lien a expiré")
    
    dossier = await db.dossiers.find_one({"id": link["dossier_id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    # Filter pieces if piece_ids specified
    query = {"dossier_id": link["dossier_id"]}
    if link.get("piece_ids") is not None and len(link.get("piece_ids", [])) > 0:
        query["id"] = {"$in": link["piece_ids"]}
    
    pieces = await db.pieces.find(
        query,
        {"_id": 0, "extracted_text": 0}
    ).sort("numero", 1).to_list(1000)
    
    # Build chronology for shared view
    chronology = []
    for p in pieces:
        if p.get("status") == "pret" and p.get("validated_data"):
            chronology.append({
                "numero": p["numero"],
                "date": p["validated_data"].get("date_document"),
                "titre": p["validated_data"].get("titre", p["original_filename"]),
                "type_piece": p["validated_data"].get("type_piece"),
                "resume": p["validated_data"]
            })
    chronology.sort(key=lambda x: x["date"] or "9999-99-99")
    
    return {
        "dossier": {
            "id": dossier["id"],
            "title": dossier["title"],
            "description": dossier["description"]
        },
        "pieces": [PieceResponse(**p) for p in pieces],
        "chronology": chronology,
        "synthesis": _compute_synthesis(pieces),
    }


def _compute_synthesis(pieces: list) -> dict:
    """Helper to compute synthesis stats from a list of pieces (used by shared + private endpoints)."""
    theme_counts: dict = {}
    subject_counts: dict = {}
    nature_counts: dict = {}
    subdomain_counts: dict = {}  # {sous_domaine: count}
    source_by_domain: dict = {}  # {domaine: {"PRO": n, "PRIVÉ": n}}
    analyzed = 0
    for p in pieces:
        c = _piece_classification(p)
        if not (c["tags_thematiques"] or c["sujets_concernes"] or c["nature_document"]):
            continue
        analyzed += 1
        for t in c["tags_thematiques"]:
            theme_counts[t] = theme_counts.get(t, 0) + 1
        for s in c["sujets_concernes"]:
            subject_counts[s] = subject_counts.get(s, 0) + 1
        if c["nature_document"]:
            nature_counts[c["nature_document"]] = nature_counts.get(c["nature_document"], 0) + 1
        if c["sous_domaine"]:
            subdomain_counts[c["sous_domaine"]] = subdomain_counts.get(c["sous_domaine"], 0) + 1
        # Ratio PRO/PRIVÉ par domaine
        src = c["source_qualifiee"]
        domain = c["tags_thematiques"][0] if c["tags_thematiques"] else None
        if domain and src in ("PRO", "PRIVÉ"):
            source_by_domain.setdefault(domain, {"PRO": 0, "PRIVÉ": 0})[src] += 1
    themes_sorted = [{"key": k, "count": v} for k, v in sorted(theme_counts.items(), key=lambda x: -x[1])]
    subjects_sorted = [{"key": k, "count": v} for k, v in sorted(subject_counts.items(), key=lambda x: -x[1])]
    natures_sorted = [{"key": k, "count": v} for k, v in sorted(nature_counts.items(), key=lambda x: -x[1])]
    subdomains_sorted = [{"key": k, "count": v} for k, v in sorted(subdomain_counts.items(), key=lambda x: -x[1])]
    hints = []
    for entry in themes_sorted[:3]:
        for h in THEME_HINTS.get(entry["key"], []):
            hints.append({"theme": entry["key"], "suggestion": h})
    return {
        "total_pieces": len(pieces),
        "pieces_classifiees": analyzed,
        "themes": themes_sorted,
        "sous_domaines": subdomains_sorted,
        "sujets": subjects_sorted,
        "natures": natures_sorted,
        "source_by_domain": source_by_domain,
        "hints": hints,
    }

@api_router.get("/shared/{token}/piece/{piece_id}/file")
async def get_shared_piece_file(token: str, piece_id: str):
    link = await db.share_links.find_one({"token": token}, {"_id": 0})
    if not link:
        raise HTTPException(status_code=404, detail="Lien non trouvé")
    
    # Check if revoked
    if link.get("revoked"):
        raise HTTPException(status_code=410, detail="Ce lien a été révoqué par son propriétaire")
    
    expires = datetime.fromisoformat(link["expires_at"])
    if datetime.now(timezone.utc) > expires:
        raise HTTPException(status_code=410, detail="Link expired")
    
    piece = await db.pieces.find_one({"id": piece_id, "dossier_id": link["dossier_id"]}, {"_id": 0})
    if not piece:
        raise HTTPException(status_code=404, detail="Piece not found")
    
    try:
        content = await storage.get_file(piece["filename"])
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="File not found")
    
    return Response(
        content=content,
        media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{piece["original_filename"]}"'}
    )

@api_router.get("/shared/{token}/export/pdf")
async def get_shared_chronology_pdf(token: str):
    """Public endpoint - Download chronology PDF"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    
    link = await db.share_links.find_one({"token": token}, {"_id": 0})
    if not link:
        raise HTTPException(status_code=404, detail="Link not found")
    
    expires = datetime.fromisoformat(link["expires_at"])
    if datetime.now(timezone.utc) > expires:
        raise HTTPException(status_code=410, detail="Link expired")
    
    dossier = await db.dossiers.find_one({"id": link["dossier_id"]}, {"_id": 0})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    pieces = await db.pieces.find({"dossier_id": link["dossier_id"], "status": "pret"}, {"_id": 0}).to_list(1000)
    
    chronology = []
    for p in pieces:
        if p.get("validated_data"):
            chronology.append({
                "numero": p["numero"],
                "date": p["validated_data"].get("date_document"),
                "titre": p["validated_data"].get("titre", p["original_filename"]),
                "resume": p["validated_data"]
            })
    chronology.sort(key=lambda x: x["date"] or "9999-99-99")
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Title_Custom', fontName='Helvetica-Bold', fontSize=16, alignment=TA_CENTER, spaceAfter=20))
    styles.add(ParagraphStyle(name='Subtitle', fontName='Helvetica', fontSize=10, alignment=TA_CENTER, textColor=colors.grey, spaceAfter=30))
    styles.add(ParagraphStyle(name='Entry_Date', fontName='Helvetica-Bold', fontSize=11, textColor=colors.HexColor('#0369A1')))
    styles.add(ParagraphStyle(name='Entry_Title', fontName='Helvetica-Bold', fontSize=10, spaceAfter=5))
    styles.add(ParagraphStyle(name='Entry_Body', fontName='Helvetica', fontSize=9, alignment=TA_JUSTIFY, spaceAfter=3))
    styles.add(ParagraphStyle(name='Entry_Ref', fontName='Helvetica-Oblique', fontSize=8, textColor=colors.grey, spaceAfter=15))
    
    elements = []
    elements.append(Paragraph("CHRONOLOGIE DES FAITS", styles['Title_Custom']))
    elements.append(Paragraph(f"Dossier : {dossier['title']}", styles['Subtitle']))
    elements.append(Spacer(1, 20))
    
    for entry in chronology:
        date_str = format_date_fr(entry["date"])
        elements.append(Paragraph(f"• {date_str}", styles['Entry_Date']))
        elements.append(Paragraph(entry["titre"], styles['Entry_Title']))
        
        resume = entry["resume"]
        if resume.get("resume_qui"):
            elements.append(Paragraph(f"<b>Qui :</b> {resume['resume_qui']}", styles['Entry_Body']))
        if resume.get("resume_quoi"):
            elements.append(Paragraph(f"<b>Quoi :</b> {resume['resume_quoi']}", styles['Entry_Body']))
        if resume.get("resume_element_cle"):
            elements.append(Paragraph(f"<b>Élément clé :</b> {resume['resume_element_cle']}", styles['Entry_Body']))
        
        elements.append(Paragraph(f"Référence : Pièce {entry['numero']}", styles['Entry_Ref']))
    
    doc.build(elements)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=chronologie.pdf"}
    )

# ===================== ROOT ROUTES =====================

# ===================== ACCOUNT DELETION =====================

class AccountDeleteRequest(BaseModel):
    immediate: bool = False  # If True, delete immediately. If False, schedule for 7 days

@api_router.delete("/account")
async def delete_account(user: dict = Depends(get_current_user), immediate: bool = False):
    """
    Delete user account.
    - immediate=False (default): Schedule deletion in 7 days (can be cancelled by logging in)
    - immediate=True: Delete immediately and permanently
    """
    user_id = user["id"]
    
    if not immediate:
        # Schedule deletion for 7 days from now
        deletion_date = (datetime.now(timezone.utc) + timedelta(days=7)).isoformat()
        await db.users.update_one(
            {"id": user_id},
            {"$set": {
                "scheduled_deletion": deletion_date,
                "updated_at": datetime.now(timezone.utc).isoformat()
            }}
        )
        logger.info(f"Account deletion scheduled: user={user_id}, date={deletion_date}")
        return {
            "message": "Suppression programmée dans 7 jours. Reconnectez-vous pour annuler.",
            "scheduled_deletion": deletion_date,
            "immediate": False
        }
    
    # Immediate deletion
    # Get all user's dossiers
    dossiers = await db.dossiers.find({"user_id": user_id}).to_list(1000)
    dossier_ids = [d["id"] for d in dossiers]
    
    # Delete all files for all pieces using storage abstraction
    pieces = await db.pieces.find({"dossier_id": {"$in": dossier_ids}}).to_list(10000)
    files_deleted = 0
    for piece in pieces:
        deleted = await storage.delete_file(piece["filename"])
        if deleted:
            files_deleted += 1
    
    # Delete all pieces
    pieces_result = await db.pieces.delete_many({"dossier_id": {"$in": dossier_ids}})
    
    # Delete all share links
    links_result = await db.share_links.delete_many({"dossier_id": {"$in": dossier_ids}})
    
    # Delete all dossiers
    dossiers_result = await db.dossiers.delete_many({"user_id": user_id})
    
    # Delete payment transactions
    payments_result = await db.payment_transactions.delete_many({"user_id": user_id})
    
    # Finally, delete the user
    await db.users.delete_one({"id": user_id})
    
    logger.info(f"Account deleted immediately: user={user_id}, dossiers={dossiers_result.deleted_count}, pieces={pieces_result.deleted_count}, files={files_deleted}")
    
    return {
        "message": "Compte supprimé définitivement",
        "immediate": True,
        "deleted": {
            "dossiers": dossiers_result.deleted_count,
            "pieces": pieces_result.deleted_count,
            "files": files_deleted,
            "share_links": links_result.deleted_count,
            "payment_transactions": payments_result.deleted_count
        }
    }

@api_router.post("/account/cancel-deletion")
async def cancel_account_deletion(user: dict = Depends(get_current_user)):
    """Cancel scheduled account deletion"""
    await db.users.update_one(
        {"id": user["id"]},
        {"$unset": {"scheduled_deletion": ""}}
    )
    return {"message": "Suppression annulée", "cancelled": True}


# ===================== SHARE LINK MANAGEMENT =====================

@api_router.get("/dossiers/{dossier_id}/share-links")
async def list_share_links(dossier_id: str, user: dict = Depends(get_current_user)):
    """List all share links for a dossier"""
    dossier = await db.dossiers.find_one({"id": dossier_id, "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=404, detail="Dossier not found")
    
    links = await db.share_links.find({"dossier_id": dossier_id}, {"_id": 0}).to_list(100)
    
    now = datetime.now(timezone.utc)
    result = []
    for link in links:
        expires = datetime.fromisoformat(link["expires_at"])
        result.append({
            **link,
            "is_expired": now > expires,
            "is_revoked": link.get("revoked", False),
            "is_active": not link.get("revoked", False) and now <= expires
        })
    
    return result


@api_router.delete("/share-links/{link_id}")
async def revoke_share_link(link_id: str, user: dict = Depends(get_current_user)):
    """Revoke a share link (cannot be undone)"""
    # Find the link
    link = await db.share_links.find_one({"id": link_id}, {"_id": 0})
    if not link:
        raise HTTPException(status_code=404, detail="Lien non trouvé")
    
    # Verify ownership
    dossier = await db.dossiers.find_one({"id": link["dossier_id"], "user_id": user["id"]})
    if not dossier:
        raise HTTPException(status_code=403, detail="Accès refusé")
    
    # Mark as revoked
    await db.share_links.update_one(
        {"id": link_id},
        {"$set": {
            "revoked": True,
            "revoked_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {"message": "Lien de partage révoqué"}


# ===================== BETA ACCESS =====================

class BetaActivateRequest(BaseModel):
    code: str

@api_router.post("/beta/activate")
async def activate_beta_code(data: BetaActivateRequest, user: dict = Depends(get_current_user)):
    """Activate premium access with beta code (for associations/testers)"""
    beta_code = os.environ.get("BETA_ACCESS_CODE", "")
    
    if not beta_code or data.code != beta_code:
        raise HTTPException(status_code=403, detail="Code invalide")
    
    # Update user to premium
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {
            "plan": "premium",
            "plan_status": "active",
            "plan_source": "beta_code",
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    logger.info(f"Beta code activated for user {user['id']}")
    
    return {"activated": True, "plan": "premium"}


@api_router.get("/")
async def root():
    return {
        "message": "Dossier Juridique Intelligent API", 
        "version": "1.0.0",
        "environment": config.ENV.value
    }

@api_router.get("/health")
async def health():
    return {
        "status": "healthy", 
        "max_file_size_mb": config.MAX_FILE_SIZE_MB,
        "environment": config.ENV.value,
        "stripe_configured": config.is_stripe_configured,
        "s3_configured": config.is_s3_configured,
        "storage_backend": config.STORAGE_BACKEND.value
    }

@api_router.post("/admin/migrate-storage")
async def migrate_local_to_gridfs(user: dict = Depends(get_current_user)):
    """Migrate files from local storage to GridFS (admin only)"""
    if user.get("plan") != "premium":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    if config.STORAGE_BACKEND.value != "gridfs":
        return {"message": "Storage backend is not gridfs", "migrated": 0}
    
    from motor.motor_asyncio import AsyncIOMotorGridFSBucket
    bucket = AsyncIOMotorGridFSBucket(db, bucket_name="file_storage")
    
    upload_dir = config.UPLOAD_DIR
    if not upload_dir.exists():
        return {"message": "No uploads directory", "migrated": 0}
    
    pieces = await db.pieces.find({}, {"filename": 1, "_id": 0}).to_list(100000)
    referenced = {p["filename"] for p in pieces}
    
    migrated = 0
    already = 0
    errors = 0
    
    for filepath in upload_dir.iterdir():
        if filepath.is_dir() or filepath.name not in referenced:
            continue
        
        existing = await bucket.find({"filename": filepath.name}).to_list(1)
        if existing:
            already += 1
            continue
        
        try:
            with open(filepath, 'rb') as f:
                content = f.read()
            await bucket.upload_from_stream(filepath.name, content)
            migrated += 1
        except Exception as e:
            logger.error(f"Migration error for {filepath.name}: {e}")
            errors += 1
    
    return {"migrated": migrated, "already_existed": already, "errors": errors}

@api_router.get("/migrate-to-emergent-storage")
async def migrate_gridfs_to_emergent():
    """Migrate files from GridFS to Emergent Object Storage. Call once after switching backends."""
    from motor.motor_asyncio import AsyncIOMotorGridFSBucket
    
    bucket = AsyncIOMotorGridFSBucket(db, bucket_name="file_storage")
    
    # Get all referenced filenames
    pieces = await db.pieces.find({}, {"filename": 1, "_id": 0}).to_list(100000)
    referenced = {p["filename"] for p in pieces}
    
    # Check what's already in Emergent storage
    migrated = 0
    already_exists = 0
    not_in_gridfs = 0
    errors = 0
    
    for filename in referenced:
        # Check if already in Emergent storage
        if await storage.file_exists(filename):
            already_exists += 1
            continue
        
        # Try to get from GridFS
        try:
            stream = await bucket.open_download_stream_by_name(filename)
            content = await stream.read()
            
            # Upload to Emergent Object Storage
            await storage.save_file(content, filename)
            migrated += 1
            
            if migrated % 5 == 0:
                logger.info(f"Migrated {migrated} files to Emergent storage...")
                
        except Exception as e:
            # File not in GridFS either
            not_in_gridfs += 1
    
    return {
        "total_referenced": len(referenced),
        "migrated": migrated,
        "already_in_emergent": already_exists,
        "not_in_gridfs": not_in_gridfs,
        "errors": errors,
        "status": "OK" if not_in_gridfs == 0 else "PARTIAL"
    }

# Include router
register_admin_routes(api_router, db, get_current_user)
app.include_router(api_router)

# Root-level health check for deployment systems (outside /api prefix)
@app.get("/health")
async def root_health():
    return {"status": "healthy"}

# Public stats (no auth) — for landing page counter
# Defined on `app` directly because it's added after include_router
@app.get("/api/public/stats")
async def public_stats():
    """Public statistics for the landing page (no auth required)."""
    try:
        total_dossiers = await db.dossiers.count_documents({})
        total_users = await db.users.count_documents({})
        return {"total_dossiers": total_dossiers, "total_users": total_users}
    except Exception as e:
        logger.error(f"public_stats error: {e}")
        return {"total_dossiers": 0, "total_users": 0}

# Security middlewares (order matters - added last = executed first)
app.add_middleware(AccessLogMiddleware)
app.add_middleware(ErrorHandlingMiddleware)
app.add_middleware(SecurityHeadersMiddleware)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=config.CORS_ORIGINS,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_db_indexes():
    """Create indexes for performance and security"""
    try:
        # Test MongoDB connection
        await client.admin.command('ping')
        logger.info("MongoDB connected")
        
        # User indexes
        await db.users.create_index("email", unique=True, name="idx_user_email")
        
        # Dossier indexes
        await db.dossiers.create_index("user_id", name="idx_dossier_user")
        
        # Piece indexes
        await db.pieces.create_index(
            [("dossier_id", 1), ("file_hash", 1), ("file_size", 1)],
            unique=False,
            name="idx_duplicate_detection"
        )
        await db.pieces.create_index([("dossier_id", 1), ("numero", 1)], name="idx_dossier_pieces")
        await db.pieces.create_index([("dossier_id", 1), ("analysis_status", 1)], name="idx_analysis_queue")
        await db.pieces.create_index([("dossier_id", 1), ("status", 1)], name="idx_piece_status")
        await db.pieces.create_index([("dossier_id", 1), ("file_type", 1)], name="idx_piece_type")
        
        # Share link indexes
        await db.share_links.create_index("token", unique=True, name="idx_share_token")
        await db.share_links.create_index("dossier_id", name="idx_share_dossier")
        await db.share_links.create_index("expires_at", name="idx_share_expiry")
        
        # Share access logs
        await db.share_access_logs.create_index("share_token", name="idx_access_log_token")
        await db.share_access_logs.create_index("timestamp", name="idx_access_log_time")
        
        # Promo codes
        await db.promo_codes.create_index("code", unique=True, name="idx_promo_code")
        
        logger.info("MongoDB indexes created successfully")
    except Exception as e:
        logger.warning(f"Index creation warning (may already exist): {e}")

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
